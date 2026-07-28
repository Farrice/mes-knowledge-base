#!/usr/bin/env python3
"""
md_to_docx.py — Convert Markdown files to formatted Google-Docs-ready .docx files.

Usage:
    python execution/md_to_docx.py <input.md> [output.docx]
    python execution/md_to_docx.py <folder/> [output_folder/]    # Batch mode
    python execution/md_to_docx.py <folder/> --upload             # Convert + upload to Drive (requires gws auth)

Features:
    - Page breaks at --- dividers and before ## headers
    - Proper heading hierarchy (H1, H2, H3)
    - Bold, italic, inline code
    - Tables with borders
    - Blockquotes with left indent
    - Bullet lists and numbered lists
    - Clean margins and fonts suitable for Google Docs import
    - Checkbox items rendered as bullet points

Created: 2026-04-12
"""

import sys
import os
import re
import glob
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_styled_doc():
    """Create a document with clean styles suitable for Google Docs."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(6)

    # Style Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2a, 0x2a, 0x2a)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Style Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Arial'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x3a, 0x3a, 0x3a)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    return doc


def add_formatted_text(paragraph, text):
    """Parse inline markdown (bold, italic, code) and add runs to paragraph."""
    # Pattern to match **bold**, *italic*, `code`, and plain text
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|[^*`]+)'
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
        else:
            paragraph.add_run(part)


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # Skip separator rows (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Style the table with borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
        tbl_pr.append(borders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): 'CCCCCC'
        })
        borders.append(border)

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                p.style = doc.styles['Normal']
                add_formatted_text(p, cell_text)
                if i == 0:  # Header row
                    for run in p.runs:
                        run.bold = True

    # Add spacing after table
    doc.add_paragraph('')


def convert_md_to_docx(md_path, docx_path=None):
    """Convert a markdown file to a formatted .docx file."""
    md_path = Path(md_path)
    if docx_path is None:
        docx_path = md_path.with_suffix('.docx')
    else:
        docx_path = Path(docx_path)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = create_styled_doc()

    i = 0
    in_blockquote = False
    in_code_block = False
    first_h1 = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue

        # Empty lines
        if not stripped:
            in_blockquote = False
            i += 1
            continue

        # Horizontal rules → page break
        if stripped in ('---', '***', '___') and len(stripped) >= 3:
            doc.add_page_break()
            i += 1
            continue

        # H1 — page break before (except first)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            text = stripped[2:].strip()
            p = doc.add_heading(text, level=1)
            i += 1
            continue

        # H2 — page break before
        if stripped.startswith('## '):
            doc.add_page_break()
            text = stripped[3:].strip()
            p = doc.add_heading(text, level=2)
            i += 1
            continue

        # H3
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            p = doc.add_heading(text, level=3)
            i += 1
            continue

        # H4
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(10)
            i += 1
            continue

        # Tables
        if stripped.startswith('|') and '|' in stripped[1:]:
            rows, end_idx = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            i = end_idx
            continue

        # Blockquotes
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            if text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                add_formatted_text(p, text)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Checkbox items
        if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
            checked = stripped.startswith('- [x]')
            text = stripped[5:].strip()
            marker = '✓' if checked else '☐'
            p = doc.add_paragraph(f'{marker}  ', style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    doc.save(str(docx_path))
    return docx_path


def batch_convert(input_dir, output_dir=None):
    """Convert all .md files in a directory to .docx."""
    input_dir = Path(input_dir)
    if output_dir is None:
        output_dir = input_dir / 'docx'
    else:
        output_dir = Path(output_dir)

    output_dir.mkdir(parents=True, exist_ok=True)

    md_files = sorted(input_dir.glob('*.md'))
    converted = []

    for md_file in md_files:
        docx_path = output_dir / md_file.with_suffix('.docx').name
        convert_md_to_docx(md_file, docx_path)
        converted.append(docx_path)
        print(f'  ✓ {md_file.name} → {docx_path.name}')

    return converted


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage:')
        print('  python execution/md_to_docx.py <file.md> [output.docx]')
        print('  python execution/md_to_docx.py <folder/> [output_folder/]')
        sys.exit(1)

    input_path = Path(sys.argv[1])

    if input_path.is_dir():
        output_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else None
        files = batch_convert(input_path, output_dir)
        print(f'\nConverted {len(files)} files.')
    elif input_path.is_file():
        output_path = sys.argv[2] if len(sys.argv) > 2 else None
        result = convert_md_to_docx(input_path, output_path)
        print(f'  ✓ {input_path.name} → {result.name}')
    else:
        print(f'Error: {input_path} not found')
        sys.exit(1)
