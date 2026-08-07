from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path("/Users/farricecain/Downloads/TrendScale_Master_Brief_Template (1).docx")


def set_cell(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)


def set_para(paragraph, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_row(table, values: list[str]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def replace_table_rows(table, rows: list[list[str]]) -> None:
    # Keep the template header row and replace all body rows with fresh rows.
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for row in rows:
        add_row(table, row)


def build_docx(brief: dict) -> Path:
    doc = Document(TEMPLATE)

    paragraph_values = {
        1: "Client-facing production brief. v3 script-quality pass for founder review and editor handoff.",
        3: f"Strategist: {brief['strategist']}",
        4: f"Client: {brief['client']}",
        5: f"PDP: {brief['pdp']}",
        6: f"Media buyer notes: {brief['media_buyer_notes']}",
        7: f"Concept type: {brief['concept_type']}",
        8: f"Iterated from: {brief['iterated_from']}",
        9: f"Avatar: {brief['avatar']}",
        10: f"Awareness level: {brief['awareness']}",
        11: f"Hypothesis: {brief['hypothesis']}",
        13: f"Concept inspo: {brief['concept_inspo']}",
        14: f"Cast type: {brief['cast_type']}",
        15: f"AI avatar look: {brief['ai_avatar']}",
        16: f"Pacing: {brief['pacing']}",
        17: f"Main emotion: {brief['main_emotion']}",
        18: f"Voiceover: {brief['voiceover']}",
        19: f"Music: {brief['music']}",
        20: f"Caption style: {brief['caption_style']}",
        21: f"Overall direction: {brief['overall_direction']}",
        22: f"Editing notes: {brief['editing_notes']}",
        24: f"Format: {brief['format']}",
        26: f"Note: {brief['note']}",
    }
    for index, text in paragraph_values.items():
        set_para(doc.paragraphs[index], text)

    replace_table_rows(doc.tables[0], brief["script_rows"])

    out_path = ROOT / brief["filename"]
    doc.save(out_path)
    return out_path


JCKED = {
    "filename": "TrendScale_JCKED_Production_Brief_v3.docx",
    "strategist": "Farrice Cain",
    "client": "JCKED - Liquid L-Carnitine 4000mg",
    "pdp": "https://jcked.com/products/liquid-l-carnitine-4000mg-of1",
    "media_buyer_notes": (
        "Control test: dose-gap label comparison. Lead with prior-bottle doubt, "
        "then show why 4,000mg liquid L-carnitine deserves a different label read. "
        "Use competitor dose contrast only when specific labels are verified."
    ),
    "concept_type": "Iteration",
    "iterated_from": "The Locked Vault control concept.",
    "avatar": (
        "Male, 35 to 50. He trains, tracks, and has already bought supplements. "
        "He is skeptical of fat-loss ads, but he will inspect a clear dose and "
        "mechanism argument."
    ),
    "awareness": "Problem aware to Solution aware",
    "hypothesis": (
        "Cold traffic will stop when the ad makes an old product failure feel "
        "unresolved. Start with the half-used L-carnitine bottle, turn on dose, "
        "and make the click a label comparison."
    ),
    "concept_inspo": (
        "Half-used bottle, cabinet mirror, locked vault, amber key, mitochondria "
        "gate, clean dose cards, real product on counter."
    ),
    "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
    "ai_avatar": (
        "Male, 38 to 45, understated, clean but lived-in. Neutral kitchen or study. "
        "Simple dark overshirt or crewneck. Calm strategist energy, no doctor coat."
    ),
    "pacing": "Hard first 3 seconds, fast dose reveal, measured mechanism, quiet product hold.",
    "main_emotion": "Vindication. The viewer should feel the old test may have been incomplete.",
    "voiceover": "Plain, clipped, and specific. Short sentences. No hype cadence.",
    "music": "Sparse low pulse at minus 18dB. Drop under the final label line.",
    "caption_style": "Manual burned-in captions. Six words per frame maximum. Amber only for key, gate, and dose-card moments.",
    "overall_direction": (
        "Use one hook, then run the body rows in order. Keep the locked-vault "
        "visual system, but make the spoken copy about dose, transport, and label "
        "inspection."
    ),
    "editing_notes": (
        "Open on a real object, not a body. Use cabinet, label, dose card, vault, "
        "and mitochondria visuals. Avoid body photos, scale imagery, before-after "
        "comparisons, doctor visuals, panic urgency, and template effects."
    ),
    "format": "Video, 9:16, 30 seconds. Optional 6-card static carousel using the dose-contrast angle.",
    "note": (
        "Structure-function language only. Final edit must match the PDP and label. "
        "Avoid hard fat-loss promises, disease claims, personal-attribute callouts, "
        "and unverified competitor claims. Source anchors: NIH ODS Carnitine "
        "Health Professional Fact Sheet, https://ods.od.nih.gov/factsheets/"
        "Carnitine-HealthProfessional/; Linus Pauling Institute L-Carnitine, "
        "https://lpi.oregonstate.edu/mic/dietary-factors/L-carnitine."
    ),
    "script_rows": [
        [
            "Half-used L-carnitine bottle in a bathroom cabinet. Thumb turns the label toward camera.",
            "Hook 1",
            "Before writing off L-carnitine, check the dose.",
            "CHECK THE DOSE",
        ],
        [
            "Two clean cards: 500mg on the left, JCKED 4,000mg on the right. No competitor logo.",
            "Hook 2",
            "A 500mg label and a 4,000mg label are two different tests.",
            "TWO DIFFERENT TESTS",
        ],
        [
            "Locked amber vault inside a stylized cell. A key stops just short of the gate.",
            "Hook 3",
            "Stored fuel still needs transport. That is the part most ads skip.",
            "FUEL NEEDS TRANSPORT",
        ],
        [
            "Cinematic mitochondria gate. L-carnitine tag moves with a long-chain fatty acid.",
            "Mechanism",
            "L-carnitine helps move long-chain fatty acids into mitochondria. That is where they can be used for energy.",
            "TRANSPORT BEFORE FUEL",
        ],
        [
            "Dose card returns. Prior bottle fades back. JCKED 4,000mg stays sharp.",
            "Dose turn",
            "So the question changes. Did the ingredient fail, or did the serving never get a fair test?",
            "DID THE DOSE HOLD UP?",
        ],
        [
            "Real JCKED bottle on a counter. Amber key line points to the 4,000mg claim.",
            "Product reveal",
            "JCKED puts 4,000mg liquid L-carnitine on the label. That is the number to look for.",
            "4,000MG LIQUID L-CARNITINE",
        ],
        [
            "Bottle hold. Quick label comparison card. CTA appears after the voice lands.",
            "CTA",
            "Read the label. Compare the serving. Then decide which test you actually ran.",
            "COMPARE THE SERVING",
        ],
    ],
}


PURAVITA = {
    "filename": "TrendScale_Puravita_Production_Brief_v3.docx",
    "strategist": "Farrice Cain",
    "client": "Puravita - 12-Form Magnesium Complex",
    "pdp": "https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex",
    "media_buyer_notes": (
        "Control test: hidden-signal metaphor. Lead with the phone or lab-report "
        "mismatch, then move to 300+ enzyme systems, the serum-magnesium caveat, "
        "form-dependent absorption, and Puravita's 12-form formula."
    ),
    "concept_type": "Iteration",
    "iterated_from": "The Battery You Can't See control concept.",
    "avatar": (
        "Men 38 to 62, with a 45 to 55 sweet spot. He has normal routines and "
        "may have normal labs, but he is trying to explain low-signal drag without "
        "turning it into a diagnosis."
    ),
    "awareness": "Unaware to Problem aware",
    "hypothesis": (
        "Cold traffic will stop for a signal story before a magnesium story. The "
        "first image is a familiar warning system: battery or lab result. Then "
        "show why magnesium is hard to read from a single surface signal."
    ),
    "concept_inspo": (
        "Phone at 5 percent, normal lab panel, kitchen morning, soft body-map glow, "
        "serum window card, form cards, quiet product reveal."
    ),
    "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
    "ai_avatar": (
        "Male, 45 to 50, understated, slight gray at temples, simple sweater or "
        "overshirt, warm kitchen light. Calm and observant."
    ),
    "pacing": "Slow first beat, sharp proof reveal, clean form explanation, quiet product hold.",
    "main_emotion": "Private recognition. The viewer should feel that the ad names a missing signal without diagnosing him.",
    "voiceover": "Quiet, grounded, and spare. Keep every sentence easy to caption.",
    "music": "Sparse piano or low ambient bed at minus 18dB. Silence under the final label line.",
    "caption_style": "Manual burned-in captions. Six words per frame maximum. Sage for serum/form callouts.",
    "overall_direction": (
        "Use one hook, then run the body rows in order. The first half is an "
        "invisible-signal story. The second half makes the 12-form label worth "
        "reading."
    ),
    "editing_notes": (
        "Use hand-only phone b-roll and clean document-style lab visuals. Use a "
        "soft body-map glow, not literal anatomy. Avoid exhausted-person stock, "
        "doctor coats, disease language, diagnosis, urgency text, and influencer "
        "quote cards unless cleared by the client."
    ),
    "format": "Video, 9:16, 35 to 45 seconds. Optional 4:5 static thumbstop using the phone or lab-report hook.",
    "note": (
        "Use label/PDP-approved structure-function claims only. Correct proof base: "
        "magnesium supports more than 300 enzyme systems; less than 1% of total "
        "body magnesium is in serum; supplement forms can differ in absorption. "
        "Avoid depletion certainty, sleep cures, fatigue claims, and unapproved "
        "named-influencer usage. Source anchor: NIH ODS Magnesium Health "
        "Professional Fact Sheet, https://ods.od.nih.gov/factsheets/"
        "Magnesium-HealthProfessional/."
    ),
    "script_rows": [
        [
            "Phone at 5 percent on a nightstand. Dawn light. Hand only, no face.",
            "Hook 1",
            "A phone shows five percent. The body gives no battery icon.",
            "NO BATTERY ICON",
        ],
        [
            "Clean lab report slides into frame. The word normal is visible. Coffee sits untouched.",
            "Hook 2",
            "The report says normal. The morning tells a messier story.",
            "NORMAL ON PAPER",
        ],
        [
            "Common sleep-ad visuals flicker off. A soft body-map glow replaces them.",
            "Hook 3",
            "Most magnesium ads chase sleep. This one starts earlier: the signal.",
            "START WITH THE SIGNAL",
        ],
        [
            "Body map lights up nerves, muscle, and energy icons. Keep it editorial, not medical.",
            "Mechanism",
            "Magnesium supports more than 300 enzyme systems. Muscle. Nerves. Energy metabolism.",
            "300+ ENZYME SYSTEMS",
        ],
        [
            "Bloodstream label shows less than 1%. Camera pulls back to bone, muscle, and soft tissue.",
            "Proof",
            "Less than 1% of total body magnesium is in serum. That makes bloodwork a narrow window.",
            "<1% IN SERUM",
        ],
        [
            "Formula board. Twelve form cards appear one by one, clean and readable.",
            "Form turn",
            "Forms matter because magnesium supplements do not absorb the same way. Puravita makes the form list the story.",
            "FORM MATTERS",
        ],
        [
            "Real Puravita bottle. Twelve form cards settle around the label.",
            "Product reveal",
            "Puravita brings 12 magnesium forms into one formula. That is why the label is worth reading.",
            "12-FORM MAGNESIUM COMPLEX",
        ],
        [
            "Bottle and supplement-facts panel. CTA appears after the final sentence.",
            "CTA",
            "Read the form list. Check the label. Learn what each form is doing.",
            "READ THE FORMULA",
        ],
    ],
}


def docx_to_markdown(path: Path) -> str:
    doc = Document(path)
    lines = [f"# {path.stem}", ""]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
            lines.append("")
    for table in doc.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        header = rows[0]
        lines.append("Script table")
        lines.append("")
        for index, row in enumerate(rows[1:], start=1):
            lines.append(f"Row {index}")
            for label, value in zip(header, row):
                lines.append(f"{label}: {value}")
            lines.append("")
        lines.append("")
    return "\n".join(lines)


def sentence_lengths(text: str) -> list[int]:
    parts = [p.strip() for p in re.split(r"[.!?]+", text) if p.strip()]
    return [len(re.findall(r"\b[\w'<%+]+\b", p)) for p in parts]


def quality_snapshot(briefs: list[dict]) -> str:
    lines = ["# TrendScale v3 Quality Snapshot", ""]
    for brief in briefs:
        script_texts = [row[2] for row in brief["script_rows"]]
        lengths = [length for text in script_texts for length in sentence_lengths(text)]
        avg = round(sum(lengths) / len(lengths), 1)
        max_len = max(lengths)
        row_word_counts = [len(re.findall(r"\b[\w'<%+]+\b", text)) for text in script_texts]
        lines.append(f"## {brief['client']}")
        lines.append(f"- Script rows: {len(script_texts)}")
        lines.append(f"- Average sentence length: {avg} words")
        lines.append(f"- Longest sentence: {max_len} words")
        lines.append(f"- Longest script row: {max(row_word_counts)} words")
        lines.append("")
    lines.append("## Client-Facing Fixes")
    lines.append("- Rebuilt hooks as one-idea openers instead of mini strategy paragraphs.")
    lines.append("- Restored the original visual metaphors while removing certainty-heavy health claims.")
    lines.append("- Replaced pending PDP notes with verified product-page URLs.")
    lines.append("- Removed internal notes and kept media buyer notes account-facing.")
    lines.append("- Kept competitor/dose contrast in JCKED, with label-verification guardrails.")
    lines.append("")
    return "\n".join(lines)


def write_text_outputs(paths: list[Path], briefs: list[dict]) -> None:
    extract = ["# TrendScale Production Briefs v3 Text Extract", ""]
    for path in paths:
        extract.append(docx_to_markdown(path))
        extract.append("")
    (ROOT / "TrendScale_Production_Briefs_v3_Text_Extract.md").write_text("\n".join(extract), encoding="utf-8")

    (ROOT / "TrendScale_v3_Autopilot_Quality_Snapshot.md").write_text(
        quality_snapshot(briefs),
        encoding="utf-8",
    )

    send_note = """Hi [Recruiter Name],

Thanks again for the feedback. I rebuilt both scripts inside the TrendScale brief format and tightened the execution.

The revised briefs keep the original concepts and research spine, but the scripts are now easier to review and produce: three hook options, a clearer mechanism/proof/product/CTA sequence, cleaner PDP and label guardrails, and tighter on-screen text. I also removed working notes so the documents read like production briefs rather than a drafting pass.

Best,
Farrice
"""
    (ROOT / "TrendScale_v3_Recruiter_Revision_Note.md").write_text(send_note, encoding="utf-8")


def main() -> None:
    briefs = [JCKED, PURAVITA]
    paths = [build_docx(brief) for brief in briefs]
    write_text_outputs(paths, briefs)
    for path in paths:
        print(path)
    print(ROOT / "TrendScale_Production_Briefs_v3_Text_Extract.md")
    print(ROOT / "TrendScale_v3_Recruiter_Revision_Note.md")
    print(ROOT / "TrendScale_v3_Autopilot_Quality_Snapshot.md")


if __name__ == "__main__":
    main()
