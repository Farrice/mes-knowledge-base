from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document

from generate_v3_production_briefs import ROOT, TEMPLATE, docx_to_markdown, set_cell, set_para


def add_row(table, values: list[str]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def replace_table_rows(table, rows: list[list[str]]) -> None:
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for row in rows:
        add_row(table, row)


def build_docx(brief: dict) -> Path:
    doc = Document(TEMPLATE)

    paragraph_values = {
        1: "Client-facing production brief for paid social review and editor handoff.",
        3: f"Strategist: {brief['strategist']}",
        4: f"Client: {brief['client']}",
        5: f"PDP: {brief['pdp']}",
        6: f"Production note: {brief['production_note']}",
        7: f"Concept type: {brief['concept_type']}",
        8: f"Iterated from: {brief['iterated_from']}",
        9: f"Avatar: {brief['avatar']}",
        10: f"Awareness level: {brief['awareness']}",
        11: f"Hypothesis: {brief['hypothesis']}",
        13: f"Concept inspo: {brief['concept_inspo']}",
        14: f"Cast type: {brief['cast_type']}",
        15: f"AI avatar look: {brief['ai_avatar']}",
        16: f"Pacing: {brief['pacing']}",
        17: f"Main emotion: {brief['main_emotion']}",
        18: f"Voiceover: {brief['voiceover']}",
        19: f"Music: {brief['music']}",
        20: f"Caption style: {brief['caption_style']}",
        21: f"Overall direction: {brief['overall_direction']}",
        22: f"Editing notes: {brief['editing_notes']}",
        24: f"Format: {brief['format']}",
        26: f"Note: {brief['note']}",
    }

    for index, text in paragraph_values.items():
        set_para(doc.paragraphs[index], text)

    replace_table_rows(doc.tables[0], brief["script_rows"])

    out_path = ROOT / brief["filename"]
    doc.save(out_path)
    client_path = ROOT / brief["client_filename"]
    shutil.copy2(out_path, client_path)
    return out_path


JCKED = {
    "filename": "TrendScale_JCKED_Production_Brief_v6.docx",
    "client_filename": "TrendScale_JCKED_Production_Brief_FINAL.docx",
    "strategist": "Farrice Cain",
    "client": "JCKED - Liquid L-Carnitine 4000MG",
    "pdp": "https://jcked.com/products/liquid-l-carnitine-4000mg-of1",
    "production_note": (
        "Recommended control is Hook 1. The script column is exact spoken voiceover. "
        "This pass is optimized for top-of-funnel curiosity and qualified click intent. "
        "The PDP verifies the product title, 4,000mg liquid L-carnitine positioning, "
        "15mL serving, and liquid delivery. Keep competitor contrast as label "
        "comparison unless specific competitor labels are cleared."
    ),
    "concept_type": "Iteration",
    "iterated_from": "The Locked Vault control concept.",
    "avatar": (
        "Male, 35 to 50. He has a cabinet graveyard of half-used supplements and "
        "does not trust transformation language. He will stop for a label contradiction: "
        "maybe the ingredient was blamed before the serving was checked."
    ),
    "awareness": "Problem aware to Solution aware",
    "hypothesis": (
        "The vault metaphor works best when the hook names the failed prior test. "
        "Lead with the private week-three quit moment, then make one label number "
        "feel like the open loop."
    ),
    "concept_inspo": (
        "Half-used bottle, bathroom cabinet, locked vault, amber key, mitochondrial "
        "gate, serving-size cards, real product on counter."
    ),
    "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
    "ai_avatar": (
        "Male, 38 to 45, understated, clean but lived-in. Neutral kitchen or study. "
        "Simple dark overshirt or crewneck. Calm strategist energy, no doctor coat."
    ),
    "pacing": "Private behavior hook, label mystery, mechanism reveal, dose card, product hold, soft click.",
    "main_emotion": "Vindication. The viewer should feel the old bottle may have been a weak test, not a final answer.",
    "voiceover": (
        "Low, plain, and forensic. Read the script column as written. No extra "
        "explainer copy."
    ),
    "music": "Sparse low pulse at minus 18dB. Drop under the final label line.",
    "caption_style": "Manual burned-in captions. Six words per frame maximum. Amber only for key, gate, and dose-card moments.",
    "overall_direction": (
        "Use one hook, then continue through the body rows in order. The ad should "
        "feel like a cold-traffic mystery about an unfair supplement test, not a "
        "fat-loss pitch."
    ),
    "editing_notes": (
        "Use cabinet, label, dose card, vault, and mitochondria visuals. Avoid body "
        "photos, scale imagery, before-after comparisons, doctor visuals, panic "
        "urgency, and template effects. Keep the bottle real."
    ),
    "format": "Video, 9:16, 30 seconds. Optional 6-card static carousel using the serving-size contrast.",
    "note": (
        "Structure-function language only. Final edit must match the live PDP and "
        "label. Avoid hard fat-loss promises, disease claims, personal-attribute "
        "callouts, and universal competitor claims. Source anchors: NIH ODS Carnitine "
        "Health Professional Fact Sheet, https://ods.od.nih.gov/factsheets/"
        "Carnitine-HealthProfessional/; Linus Pauling Institute L-Carnitine, "
        "https://lpi.oregonstate.edu/mic/dietary-factors/L-carnitine."
    ),
    "script_rows": [
        [
            "Half-used L-carnitine bottle in a bathroom cabinet. Thumb turns the label toward camera.",
            "Hook 1",
            "He quit L-carnitine in week three. The bottle may prove the test was never fair.",
            "THE TEST WAS NEVER FAIR",
        ],
        [
            "Locked amber vault inside a stylized cell. A key stops just short of the gate.",
            "Hook 2",
            "The ingredient did not need a louder promise. It needed a fair serving.",
            "FAIR SERVING",
        ],
        [
            "Two clean cards: 500mg on the left, JCKED 4,000mg on the right. No competitor logo.",
            "Hook 3",
            "A low-dose bottle and a 4,000mg liquid serving are not the same experiment.",
            "NOT THE SAME TEST",
        ],
        [
            "Cinematic mitochondria gate. L-carnitine tag moves with a long-chain fatty acid.",
            "Mechanism",
            "Long-chain fatty acids do not just drift into mitochondria. They need a transport step.",
            "THE TRANSPORT STEP",
        ],
        [
            "Amber key reaches the vault. Gate opens just enough to show the metaphor.",
            "Vault turn",
            "L-carnitine is part of that step. That is the locked vault.",
            "THE LOCKED VAULT",
        ],
        [
            "Dose card returns. Prior bottle fades back. JCKED 4,000mg stays sharp.",
            "Dose turn",
            "That is why the label number changes the whole experiment.",
            "THE LABEL CHANGES THE TEST",
        ],
        [
            "Real JCKED bottle on a counter. Macro label hold on 4,000mg and serving size.",
            "Product reveal",
            "JCKED makes that question visible: 4,000mg liquid L-carnitine in a 15mL serving.",
            "4,000MG / 15ML",
        ],
        [
            "Bottle hold. Quick label comparison card. CTA appears after the voice lands.",
            "CTA",
            "Open the product page. Check the serving. Then decide if the old bottle got a fair shot.",
            "CHECK THE SERVING",
        ],
    ],
}


PURAVITA = {
    "filename": "TrendScale_Puravita_Production_Brief_v6.docx",
    "client_filename": "TrendScale_Puravita_Production_Brief_FINAL.docx",
    "strategist": "Farrice Cain",
    "client": "Puravita - Magnesium Complex Capsules",
    "pdp": "https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex",
    "production_note": (
        "Recommended control is Hook 1. The script column is exact spoken voiceover. "
        "This pass is optimized for top-of-funnel curiosity and qualified click intent. "
        "The PDP verifies the active Puravita Magnesium Complex Capsules listing and "
        "quantity offers. Keep the 12-form language from the original product brief "
        "and final label, without making form-specific health promises."
    ),
    "concept_type": "Iteration",
    "iterated_from": "The Battery You Can't See control concept.",
    "avatar": (
        "Men 38 to 62, with a 45 to 55 sweet spot. He has routines, labs, and enough "
        "life pressure to distrust another generic magnesium sleep ad. He needs a "
        "clean reason to inspect the label."
    ),
    "awareness": "Unaware to Problem aware",
    "hypothesis": (
        "The battery concept works best when the hook starts from the private search, "
        "not the supplement category. Name the 11pm question first, then use serum as "
        "the curiosity bridge into the 12-form label."
    ),
    "concept_inspo": (
        "Phone at 5 percent, normal lab panel, untouched coffee, soft body-map glow, "
        "serum window card, form list, quiet product reveal."
    ),
    "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
    "ai_avatar": (
        "Male, 45 to 50, understated, slight gray at temples, simple sweater or "
        "overshirt, warm kitchen light. Calm and observant."
    ),
    "pacing": "Private search hook, signal turn, proof card, form-board reveal, bottle hold.",
    "main_emotion": "Recognition without diagnosis. The viewer should feel the ad named a gap without naming him.",
    "voiceover": (
        "Quiet, restrained, and observational. Read the script column as written. "
        "No extra explainer copy."
    ),
    "music": "Sparse piano or low ambient bed at minus 18dB. Silence under the final label line.",
    "caption_style": "Manual burned-in captions. Six words per frame maximum. Sage for serum and form callouts.",
    "overall_direction": (
        "Use one hook, then continue through the body rows in order. The ad should "
        "move from private question to label inspection, without diagnosing the viewer."
    ),
    "editing_notes": (
        "Use hand-only phone b-roll and clean document-style lab visuals. Use a "
        "soft body-map glow, not literal anatomy. Avoid exhausted-person stock, "
        "doctor coats, disease language, diagnosis, urgency text, and influencer "
        "quote cards unless cleared by the client."
    ),
    "format": "Video, 9:16, 35 to 45 seconds. Optional 4:5 static thumbstop using the phone or lab-report hook.",
    "note": (
        "Use label/PDP-approved structure-function claims only. Correct proof base: "
        "magnesium supports more than 300 enzyme systems; less than 1% of total "
        "body magnesium is in serum; supplement forms can differ in absorption. "
        "Avoid depletion certainty, sleep cures, fatigue claims, and unapproved "
        "named-influencer usage. Source anchor: NIH ODS Magnesium Health "
        "Professional Fact Sheet, https://ods.od.nih.gov/factsheets/"
        "Magnesium-HealthProfessional/."
    ),
    "script_rows": [
        [
            "Phone at 5 percent on a nightstand. Dawn light. Hand only, no face.",
            "Hook 1",
            "At 11:07, he is not buying magnesium. He is trying to figure out why normal still feels unfinished.",
            "TRYING TO NAME IT",
        ],
        [
            "Clean lab report slides into frame. The word normal is visible. Coffee sits untouched.",
            "Hook 2",
            "The word normal can end the search before the mineral question even starts.",
            "NORMAL CAN END THE SEARCH",
        ],
        [
            "Common sleep-ad visuals flicker off. A soft body-map glow replaces them.",
            "Hook 3",
            "Phones show 5 percent. Bodies leave clues.",
            "NO BATTERY ICON",
        ],
        [
            "Body map lights up energy, nerve, and muscle icons. Keep it editorial, not medical.",
            "Mechanism",
            "Magnesium is one clue worth inspecting. It supports more than 300 enzyme systems, including energy, nerve, and muscle function.",
            "300+ ENZYME SYSTEMS",
        ],
        [
            "Bloodstream label shows less than 1%. Camera pulls back to bone, muscle, and soft tissue.",
            "Proof",
            "Less than 1% of total body magnesium is in serum. That is a tiny window.",
            "<1% IN SERUM",
        ],
        [
            "Formula board. Twelve form cards appear one by one, clean and readable.",
            "Form turn",
            "Then the label gets strategic. Magnesium forms do not all absorb the same way.",
            "FORM CHANGES THE LABEL",
        ],
        [
            "Real Puravita bottle. Twelve form cards settle around the label.",
            "Product reveal",
            "Puravita puts 12 magnesium forms in one formula. That turns the form list into the story.",
            "12 MAGNESIUM FORMS",
        ],
        [
            "Bottle and supplement-facts panel. CTA appears after the final sentence.",
            "CTA",
            "Open the product page. Read the form list. Then compare it to the simpler magnesium bottles around it.",
            "READ THE FORM LIST",
        ],
    ],
}


def sentence_lengths(text: str) -> list[int]:
    parts = [p.strip() for p in re.split(r"[.!?]+", text) if p.strip()]
    return [len(re.findall(r"\b[\w'<%+]+\b", p)) for p in parts]


def full_script(brief: dict, hook_index: int = 0) -> str:
    rows = brief["script_rows"]
    selected_hook = rows[hook_index][2]
    body = [row[2] for row in rows[3:]]
    return " ".join([selected_hook, *body])


def quality_snapshot(briefs: list[dict]) -> str:
    lines = ["# TrendScale v6 Conversion Hook Quality Snapshot", ""]
    for brief in briefs:
        script_texts = [row[2] for row in brief["script_rows"]]
        lengths = [length for text in script_texts for length in sentence_lengths(text)]
        avg = round(sum(lengths) / len(lengths), 1)
        max_len = max(lengths)
        full = full_script(brief)
        full_word_count = len(re.findall(r"\b[\w'<%+]+\b", full))
        lines.append(f"## {brief['client']}")
        lines.append(f"- Script rows: {len(script_texts)}")
        lines.append(f"- Full control voiceover length: {full_word_count} words")
        lines.append(f"- Average sentence length: {avg} words")
        lines.append(f"- Longest sentence: {max_len} words")
        lines.append(f"- Recommended control hook: Hook 1")
        lines.append("")
    lines.append("## What changed from v5")
    lines.append("- Rebuilt Hook 1 for each product around private behavior, not explanation.")
    lines.append("- JCKED now opens with the week-three quit and the unfair-test loop.")
    lines.append("- Puravita now opens at 11:07 with the unresolved normal-still-feels-unfinished loop.")
    lines.append("- Added stronger curiosity debt while keeping proof and supplement claims bounded.")
    lines.append("- Kept the brief easy for an editor to shoot: one hook, then body rows in order.")
    lines.append("")
    return "\n".join(lines)


def red_team_receipt(briefs: list[dict]) -> str:
    return """# TrendScale v6 Conversion Hook Red-Team Receipt

## Failure mode addressed
- v5 was structurally clean and recordable, but Hook 1 still felt like a 3/10 cold-traffic hook.
- v6 raises the top-of-funnel entry by using private behavior, curiosity debt, and one unresolved mechanism question per product.

## Claim/proof ledger
| Claim | Status | Source/Proof | Brief decision |
|---|---|---|---|
| JCKED product page is live | VERIFIED | https://jcked.com/products/liquid-l-carnitine-4000mg-of1 | Kept |
| JCKED product title is Liquid L-Carnitine 4000MG | VERIFIED | Shopify product JSON | Kept |
| JCKED PDP references 4,000mg liquid L-carnitine and 15mL serving | VERIFIED | Shopify product JSON | Kept |
| L-carnitine helps transport long-chain fatty acids into mitochondria | VERIFIED | NIH ODS Carnitine and Linus Pauling Institute | Kept as structure-function mechanism |
| Puravita product page is live | VERIFIED | https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex | Kept |
| Puravita product title is Magnesium Complex Capsules | VERIFIED | Shopify product JSON | Kept |
| Puravita has 12-form magnesium positioning | PRODUCT BRIEF / LABEL-DEPENDENT | Original TrendScale source packet | Kept, with label-confirmation guardrail |
| Magnesium supports more than 300 enzyme systems | VERIFIED | NIH ODS Magnesium | Kept |
| Less than 1% of total magnesium is in serum | VERIFIED | NIH ODS Magnesium | Kept |
| Magnesium supplement forms vary in absorption | VERIFIED | NIH ODS Magnesium | Kept |

## Copy red-team
- No hard fat-loss promises.
- No disease claims.
- No depletion certainty.
- No sleep cure claims.
- No Huberman or Attia dependency.
- No universal competitor claim like "most brands underdose."
- No internal recruiter, founder, AI, or media-buyer notes.
- No body-shaming, before-after, scale, or diagnosis framing in the script.
- Hook 1 for each product creates an open loop without front-loading education.
- Product appears only after mechanism/proof earns it.

## Founder objection answers
- If the founder says the hooks needed more force: Hook 1 now starts with the private behavior the buyer recognizes before the product appears.
- If the founder worries the copy is too careful: the contrast is still there, but the risk shifted from unsupported outcome claims to label inspection and serum-window curiosity.
- If the founder asks why 12 forms matter: the ad does not over-explain individual forms; it makes the form list the reason to click.
- If the founder asks why the product appears late: the body earns the product reveal with mechanism and proof, then lands on the PDP.
- If the founder asks what to test first: run Hook 1 as control for both briefs, then test Hook 3 as the more contrastive variant.

## Verdict
SEND-READY for recruiter/founder review after final brand label/compliance review.
"""


def composition_ledger() -> str:
    return """# TrendScale v6 Expert Composition Ledger

## Owner
copy-engine, with Luke Iha copy-block sequencing as the spine.

## Accepted Specialist Pressure
| Slot | Accepted move | Evidence in v6 |
|---|---|---|
| Luke Iha vicious hooks | Private behavior first, then unresolved consequence | JCKED opens on the week-three quit; Puravita opens at 11:07 before the product pitch. |
| Stefan Georgi dopamine copy | Curiosity debt before education | JCKED withholds the serving-size answer until the dose turn; Puravita withholds the serum/form answer until after the private-search hook. |
| Dara Denney Meta ads | Top-of-funnel contrast that can survive paid traffic | Both controls lead with a recognizable failed test, then move into a label inspection reason to click. |
| Harry Dry specificity audit | Visual, falsifiable, unique details | Week three, 4,000mg, 15mL, 11:07, less than 1% in serum, and 12 magnesium forms. |
| Claim-safety gate | Bounded structure-function language | No hard fat-loss promise, no sleep cure, no depletion certainty, no disease claim, no named-influencer dependency, no universal competitor accusation. |

## Skipped Pressure
- Did not add broad persuasion frameworks that would make the brief feel overbuilt.
- Did not add media-buyer-only testing notes to the client-facing documents.
- Did not use real subagents; this was a main-thread integration pass with expert workflows loaded as standards.

## Integration Verdict
The v6 briefs are no longer just clean strategy documents. The script rows now read as recordable paid-social voiceover, with the concept, proof, product reveal, and PDP click reason joined into one sequence.
"""


def recruiter_note() -> str:
    return """Hi [Recruiter Name],

Thanks again for the direction. I made one final conversion pass on both TrendScale briefs so the hooks and script rows now read as paid-ad voiceover, not strategy notes.

The concepts are unchanged: JCKED is still The Locked Vault, and Puravita is still The Battery You Can't See. This pass strengthens the top-of-funnel hooks, keeps the research proof inside the brief, adds the live PDP context, and keeps the claim language ready for founder and production review.

Attachment order:
1. TrendScale_JCKED_Production_Brief_FINAL.docx
2. TrendScale_Puravita_Production_Brief_FINAL.docx

Best,
Farrice
"""


def write_outputs(paths: list[Path], briefs: list[dict]) -> None:
    extract = ["# TrendScale Production Briefs v6 Conversion Text Extract", ""]
    for path in paths:
        extract.append(docx_to_markdown(path))
        extract.append("")
    (ROOT / "TrendScale_Production_Briefs_v6_Conversion_Text_Extract.md").write_text(
        "\n".join(extract),
        encoding="utf-8",
    )

    (ROOT / "TrendScale_v6_Conversion_Hook_Quality_Snapshot.md").write_text(
        quality_snapshot(briefs),
        encoding="utf-8",
    )

    (ROOT / "TrendScale_v6_Conversion_RedTeam_Receipt.md").write_text(
        red_team_receipt(briefs),
        encoding="utf-8",
    )

    (ROOT / "TrendScale_v6_Expert_Composition_Ledger.md").write_text(
        composition_ledger(),
        encoding="utf-8",
    )

    (ROOT / "TrendScale_v6_Recruiter_Revision_Note.md").write_text(
        recruiter_note(),
        encoding="utf-8",
    )

    package = ROOT / "TrendScale_v6_Conversion_Send_Package.zip"
    package_items = [
        ROOT / "TrendScale_JCKED_Production_Brief_FINAL.docx",
        ROOT / "TrendScale_Puravita_Production_Brief_FINAL.docx",
        ROOT / "TrendScale_v6_Recruiter_Revision_Note.md",
    ]
    with zipfile.ZipFile(package, "w", zipfile.ZIP_DEFLATED) as archive:
        for item in package_items:
            archive.write(item, item.name)


def main() -> None:
    briefs = [JCKED, PURAVITA]
    paths = [build_docx(brief) for brief in briefs]
    write_outputs(paths, briefs)
    for path in paths:
        print(path)
    print(ROOT / "TrendScale_Production_Briefs_v6_Conversion_Text_Extract.md")
    print(ROOT / "TrendScale_v6_Recruiter_Revision_Note.md")
    print(ROOT / "TrendScale_v6_Conversion_Hook_Quality_Snapshot.md")
    print(ROOT / "TrendScale_v6_Conversion_RedTeam_Receipt.md")
    print(ROOT / "TrendScale_v6_Expert_Composition_Ledger.md")
    print(ROOT / "TrendScale_v6_Conversion_Send_Package.zip")


if __name__ == "__main__":
    main()
