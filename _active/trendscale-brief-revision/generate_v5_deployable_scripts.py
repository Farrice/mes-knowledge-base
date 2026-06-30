from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document

from generate_v3_production_briefs import ROOT, TEMPLATE, docx_to_markdown, set_cell, set_para


def add_row(table, values: list[str]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def replace_table_rows(table, rows: list[list[str]]) -> None:
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for row in rows:
        add_row(table, row)


def build_docx(brief: dict) -> Path:
    doc = Document(TEMPLATE)

    paragraph_values = {
        1: "Client-facing production brief for paid social review and editor handoff.",
        3: f"Strategist: {brief['strategist']}",
        4: f"Client: {brief['client']}",
        5: f"PDP: {brief['pdp']}",
        6: f"Production note: {brief['production_note']}",
        7: f"Concept type: {brief['concept_type']}",
        8: f"Iterated from: {brief['iterated_from']}",
        9: f"Avatar: {brief['avatar']}",
        10: f"Awareness level: {brief['awareness']}",
        11: f"Hypothesis: {brief['hypothesis']}",
        13: f"Concept inspo: {brief['concept_inspo']}",
        14: f"Cast type: {brief['cast_type']}",
        15: f"AI avatar look: {brief['ai_avatar']}",
        16: f"Pacing: {brief['pacing']}",
        17: f"Main emotion: {brief['main_emotion']}",
        18: f"Voiceover: {brief['voiceover']}",
        19: f"Music: {brief['music']}",
        20: f"Caption style: {brief['caption_style']}",
        21: f"Overall direction: {brief['overall_direction']}",
        22: f"Editing notes: {brief['editing_notes']}",
        24: f"Format: {brief['format']}",
        26: f"Note: {brief['note']}",
    }

    for index, text in paragraph_values.items():
        set_para(doc.paragraphs[index], text)

    replace_table_rows(doc.tables[0], brief["script_rows"])

    out_path = ROOT / brief["filename"]
    doc.save(out_path)
    client_path = ROOT / brief["client_filename"]
    shutil.copy2(out_path, client_path)
    return out_path


JCKED = {
    "filename": "TrendScale_JCKED_Production_Brief_v5.docx",
    "client_filename": "TrendScale_JCKED_Production_Brief_FINAL.docx",
    "strategist": "Farrice Cain",
    "client": "JCKED - Liquid L-Carnitine 4000MG",
    "pdp": "https://jcked.com/products/liquid-l-carnitine-4000mg-of1",
    "production_note": (
        "Recommended control is Hook 1. The script column is exact spoken voiceover. "
        "The PDP verifies the product title, 4,000mg liquid L-carnitine positioning, "
        "15mL serving, and liquid delivery. Keep the competitor contrast as label "
        "comparison unless specific competitor labels are cleared."
    ),
    "concept_type": "Iteration",
    "iterated_from": "The Locked Vault control concept.",
    "avatar": (
        "Male, 35 to 50. He has a cabinet graveyard of half-used supplements and "
        "does not trust transformation language. He will stop for a label contradiction: "
        "maybe the ingredient was blamed before the serving was checked."
    ),
    "awareness": "Problem aware to Solution aware",
    "hypothesis": (
        "The vault metaphor works when it reopens an old purchase. Start with the "
        "bottle he already dismissed, turn the label around, then make 4,000mg feel "
        "like the number that changes the read."
    ),
    "concept_inspo": (
        "Half-used bottle, bathroom cabinet, locked vault, amber key, mitochondrial "
        "gate, serving-size cards, real product on counter."
    ),
    "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
    "ai_avatar": (
        "Male, 38 to 45, understated, clean but lived-in. Neutral kitchen or study. "
        "Simple dark overshirt or crewneck. Calm strategist energy, no doctor coat."
    ),
    "pacing": "Object open, label turn, mechanism, dose card, product hold, soft click.",
    "main_emotion": "Vindication. The viewer should feel the old bottle may have been a weak test, not a final answer.",
    "voiceover": (
        "Low, plain, and forensic. Read the script column as written. No extra "
        "explainer copy."
    ),
    "music": "Sparse low pulse at minus 18dB. Drop under the final label line.",
    "caption_style": "Manual burned-in captions. Six words per frame maximum. Amber only for key, gate, and dose-card moments.",
    "overall_direction": (
        "Use one hook, then continue through the body rows in order. The ad should "
        "feel like a skeptical label check, not a fat-loss pitch."
    ),
    "editing_notes": (
        "Use cabinet, label, dose card, vault, and mitochondria visuals. Avoid body "
        "photos, scale imagery, before-after comparisons, doctor visuals, panic "
        "urgency, and template effects. Keep the bottle real."
    ),
    "format": "Video, 9:16, 30 seconds. Optional 6-card static carousel using the serving-size contrast.",
    "note": (
        "Structure-function language only. Final edit must match the live PDP and "
        "label. Avoid hard fat-loss promises, disease claims, personal-attribute "
        "callouts, and universal competitor claims. Source anchors: NIH ODS Carnitine "
        "Health Professional Fact Sheet, https://ods.od.nih.gov/factsheets/"
        "Carnitine-HealthProfessional/; Linus Pauling Institute L-Carnitine, "
        "https://lpi.oregonstate.edu/mic/dietary-factors/L-carnitine."
    ),
    "script_rows": [
        [
            "Half-used L-carnitine bottle in a bathroom cabinet. Thumb turns the label toward camera.",
            "Hook 1",
            "That L-carnitine bottle in the cabinet may have blamed the wrong thing. Before the ingredient takes the hit, turn the label around.",
            "TURN THE LABEL AROUND",
        ],
        [
            "Locked amber vault inside a stylized cell. A key stops just short of the gate.",
            "Hook 2",
            "Stored fuel can sit behind a locked door. L-carnitine is part of the transport step that gets it to the furnace.",
            "THE DOOR STAYS SHUT",
        ],
        [
            "Two clean cards: 500mg on the left, JCKED 4,000mg on the right. No competitor logo.",
            "Hook 3",
            "A 500mg label and a 4,000mg label are not the same test. Judge them the same, and the ingredient gets the blame.",
            "DIFFERENT TESTS",
        ],
        [
            "Cinematic mitochondria gate. L-carnitine tag moves with a long-chain fatty acid.",
            "Mechanism",
            "L-carnitine helps move long-chain fatty acids into mitochondria, where they can be used for energy.",
            "FATTY ACID TRANSPORT",
        ],
        [
            "Amber key reaches the vault. Gate opens just enough to show the metaphor.",
            "Vault turn",
            "The vault idea is simple: the key only matters if there is enough of it to test.",
            "THE LOCKED VAULT",
        ],
        [
            "Real JCKED bottle on a counter. Macro label hold on 4,000mg and serving size.",
            "Product reveal",
            "JCKED puts 4,000mg liquid L-carnitine on the label, in a 15mL serving.",
            "4,000MG / 15ML",
        ],
        [
            "Dose card returns. Prior bottle fades back. JCKED 4,000mg stays sharp.",
            "Decision turn",
            "That is the whole bet: serving first, promise second.",
            "LABEL FIRST",
        ],
        [
            "Bottle hold. Quick label comparison card. CTA appears after the voice lands.",
            "CTA",
            "Open the product page. Read the serving size. Then compare it to the bottle already sitting in the cabinet.",
            "COMPARE THE SERVING",
        ],
    ],
}


PURAVITA = {
    "filename": "TrendScale_Puravita_Production_Brief_v5.docx",
    "client_filename": "TrendScale_Puravita_Production_Brief_FINAL.docx",
    "strategist": "Farrice Cain",
    "client": "Puravita - Magnesium Complex Capsules",
    "pdp": "https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex",
    "production_note": (
        "Recommended control is Hook 1. The script column is exact spoken voiceover. "
        "The PDP verifies the active Puravita Magnesium Complex Capsules listing and "
        "quantity offers. Keep the 12-form language from the original product brief "
        "and final label, without making form-specific health promises."
    ),
    "concept_type": "Iteration",
    "iterated_from": "The Battery You Can't See control concept.",
    "avatar": (
        "Men 38 to 62, with a 45 to 55 sweet spot. He has routines, labs, and enough "
        "life pressure to distrust another generic magnesium sleep ad. He needs a "
        "clean reason to inspect the label."
    ),
    "awareness": "Unaware to Problem aware",
    "hypothesis": (
        "The battery concept works when the first beat feels like a missing dashboard, "
        "not a symptom list. The serum-window proof opens the signal. The form list "
        "gives Puravita the click."
    ),
    "concept_inspo": (
        "Phone at 5 percent, normal lab panel, untouched coffee, soft body-map glow, "
        "serum window card, form list, quiet product reveal."
    ),
    "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
    "ai_avatar": (
        "Male, 45 to 50, understated, slight gray at temples, simple sweater or "
        "overshirt, warm kitchen light. Calm and observant."
    ),
    "pacing": "Quiet first beat, signal turn, proof card, form-board reveal, bottle hold.",
    "main_emotion": "Recognition without diagnosis. The viewer should feel the ad named a gap without naming him.",
    "voiceover": (
        "Quiet, restrained, and observational. Read the script column as written. "
        "No extra explainer copy."
    ),
    "music": "Sparse piano or low ambient bed at minus 18dB. Silence under the final label line.",
    "caption_style": "Manual burned-in captions. Six words per frame maximum. Sage for serum and form callouts.",
    "overall_direction": (
        "Use one hook, then continue through the body rows in order. The ad should "
        "move from invisible signal to label inspection, without diagnosing the viewer."
    ),
    "editing_notes": (
        "Use hand-only phone b-roll and clean document-style lab visuals. Use a "
        "soft body-map glow, not literal anatomy. Avoid exhausted-person stock, "
        "doctor coats, disease language, diagnosis, urgency text, and influencer "
        "quote cards unless cleared by the client."
    ),
    "format": "Video, 9:16, 35 to 45 seconds. Optional 4:5 static thumbstop using the phone or lab-report hook.",
    "note": (
        "Use label/PDP-approved structure-function claims only. Correct proof base: "
        "magnesium supports more than 300 enzyme systems; less than 1% of total "
        "body magnesium is in serum; supplement forms can differ in absorption. "
        "Avoid depletion certainty, sleep cures, fatigue claims, and unapproved "
        "named-influencer usage. Source anchor: NIH ODS Magnesium Health "
        "Professional Fact Sheet, https://ods.od.nih.gov/factsheets/"
        "Magnesium-HealthProfessional/."
    ),
    "script_rows": [
        [
            "Phone at 5 percent on a nightstand. Dawn light. Hand only, no face.",
            "Hook 1",
            "A phone warns you at 5 percent. The body does not give you a battery icon.",
            "NO BODY BATTERY ICON",
        ],
        [
            "Clean lab report slides into frame. The word normal is visible. Coffee sits untouched.",
            "Hook 2",
            "The report says normal. That does not always tell the whole magnesium story.",
            "NORMAL ON PAPER",
        ],
        [
            "Common sleep-ad visuals flicker off. A soft body-map glow replaces them.",
            "Hook 3",
            "Most magnesium ads start with sleep. This one starts with the tiny window called serum.",
            "BEFORE SLEEP: SIGNAL",
        ],
        [
            "Body map lights up energy, nerve, and muscle icons. Keep it editorial, not medical.",
            "Mechanism",
            "Magnesium supports more than 300 enzyme systems, including energy metabolism, nerve function, and muscle function.",
            "300+ ENZYME SYSTEMS",
        ],
        [
            "Bloodstream label shows less than 1%. Camera pulls back to bone, muscle, and soft tissue.",
            "Proof",
            "Less than 1% of total body magnesium is in serum. So serum can be a narrow window into a bigger system.",
            "<1% IN SERUM",
        ],
        [
            "Formula board. Twelve form cards appear one by one, clean and readable.",
            "Form turn",
            "Then the label matters. Magnesium supplements do not all absorb the same way.",
            "FORM CHANGES THE LABEL",
        ],
        [
            "Real Puravita bottle. Twelve form cards settle around the label.",
            "Product reveal",
            "Puravita puts 12 magnesium forms in one formula. Now the form list is the story.",
            "12 MAGNESIUM FORMS",
        ],
        [
            "Bottle and supplement-facts panel. CTA appears after the final sentence.",
            "CTA",
            "Open the product page. Read the form list. Then compare it to a generic sleep capsule.",
            "READ THE FORM LIST",
        ],
    ],
}


def sentence_lengths(text: str) -> list[int]:
    parts = [p.strip() for p in re.split(r"[.!?]+", text) if p.strip()]
    return [len(re.findall(r"\b[\w'<%+]+\b", p)) for p in parts]


def full_script(brief: dict, hook_index: int = 0) -> str:
    rows = brief["script_rows"]
    selected_hook = rows[hook_index][2]
    body = [row[2] for row in rows[3:]]
    return " ".join([selected_hook, *body])


def quality_snapshot(briefs: list[dict]) -> str:
    lines = ["# TrendScale v5 Deployable Script Quality Snapshot", ""]
    for brief in briefs:
        script_texts = [row[2] for row in brief["script_rows"]]
        lengths = [length for text in script_texts for length in sentence_lengths(text)]
        avg = round(sum(lengths) / len(lengths), 1)
        max_len = max(lengths)
        full = full_script(brief)
        full_word_count = len(re.findall(r"\b[\w'<%+]+\b", full))
        lines.append(f"## {brief['client']}")
        lines.append(f"- Script rows: {len(script_texts)}")
        lines.append(f"- Full control voiceover length: {full_word_count} words")
        lines.append(f"- Average sentence length: {avg} words")
        lines.append(f"- Longest sentence: {max_len} words")
        lines.append(f"- Recommended control hook: Hook 1")
        lines.append("")
    lines.append("## What changed from v4")
    lines.append("- Rebuilt the script column as exact spoken voiceover, not production description.")
    lines.append("- Reintroduced the stronger original scenes: cabinet bottle, label turn, low battery, lab report.")
    lines.append("- Kept proof near the line it supports, without making treatment or outcome promises.")
    lines.append("- Removed internal media-buyer language and replaced it with client-facing production/PDP facts.")
    lines.append("- Kept the brief easy for an editor to shoot: one hook, then body rows in order.")
    lines.append("")
    return "\n".join(lines)


def red_team_receipt(briefs: list[dict]) -> str:
    return """# TrendScale v5 Script Red-Team Receipt

## Failure mode addressed
- v4 was safe but too caption-like. The script rows described the idea instead of sounding like paid-ad voiceover.
- v5 makes each script row recordable as spoken copy.

## Claim/proof ledger
| Claim | Status | Source/Proof | Brief decision |
|---|---|---|---|
| JCKED product page is live | VERIFIED | https://jcked.com/products/liquid-l-carnitine-4000mg-of1 | Kept |
| JCKED product title is Liquid L-Carnitine 4000MG | VERIFIED | Shopify product JSON | Kept |
| JCKED PDP references 4,000mg liquid L-carnitine and 15mL serving | VERIFIED | Shopify product JSON | Kept |
| L-carnitine helps transport long-chain fatty acids into mitochondria | VERIFIED | NIH ODS Carnitine and Linus Pauling Institute | Kept as structure-function mechanism |
| Puravita product page is live | VERIFIED | https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex | Kept |
| Puravita product title is Magnesium Complex Capsules | VERIFIED | Shopify product JSON | Kept |
| Puravita has 12-form magnesium positioning | PRODUCT BRIEF / LABEL-DEPENDENT | Original TrendScale source packet | Kept, with label-confirmation guardrail |
| Magnesium supports more than 300 enzyme systems | VERIFIED | NIH ODS Magnesium | Kept |
| Less than 1% of total magnesium is in serum | VERIFIED | NIH ODS Magnesium | Kept |
| Magnesium supplement forms vary in absorption | VERIFIED | NIH ODS Magnesium | Kept |

## Copy red-team
- No hard fat-loss promises.
- No disease claims.
- No depletion certainty.
- No sleep cure claims.
- No Huberman or Attia dependency.
- No universal competitor claim like "most brands underdose."
- No internal recruiter, founder, AI, or media-buyer notes.
- No body-shaming, before-after, scale, or diagnosis framing in the script.

## Founder objection answers
- If the founder says the hooks still need the original force: Hook 1 restores the object-led opener and gives the actor an actual first line to record.
- If the founder worries the copy is too careful: the contrast is still there, but the risk shifted from unsupported outcome claims to label inspection.
- If the founder asks why 12 forms matter: the ad does not over-explain individual forms; it makes the form list the reason to click.
- If the founder asks why the product appears late: the body earns the product reveal with mechanism and proof, then lands on the PDP.
- If the founder asks what to test first: run Hook 1 as control for both briefs, then test Hook 3 as the more contrastive variant.

## Verdict
SEND-READY for recruiter/founder review after final brand label/compliance review.
"""


def recruiter_note() -> str:
    return """Hi [Recruiter Name],

Thanks again for the direction. I made one final script pass on both TrendScale briefs so the script rows now read as exact paid-ad voiceover, not strategy notes.

The concepts are unchanged: JCKED is still The Locked Vault, and Puravita is still The Battery You Can't See. This pass tightens the hooks, keeps the research proof inside the brief, adds the live PDP context, and keeps the claim language ready for founder and production review.

Attachment order:
1. TrendScale_JCKED_Production_Brief_FINAL.docx
2. TrendScale_Puravita_Production_Brief_FINAL.docx

Best,
Farrice
"""


def write_outputs(paths: list[Path], briefs: list[dict]) -> None:
    extract = ["# TrendScale Production Briefs v5 Deployable Text Extract", ""]
    for path in paths:
        extract.append(docx_to_markdown(path))
        extract.append("")
    (ROOT / "TrendScale_Production_Briefs_v5_Deployable_Text_Extract.md").write_text(
        "\n".join(extract),
        encoding="utf-8",
    )

    (ROOT / "TrendScale_v5_Deployable_Script_Quality_Snapshot.md").write_text(
        quality_snapshot(briefs),
        encoding="utf-8",
    )

    (ROOT / "TrendScale_v5_Script_RedTeam_Receipt.md").write_text(
        red_team_receipt(briefs),
        encoding="utf-8",
    )

    (ROOT / "TrendScale_v5_Recruiter_Revision_Note.md").write_text(
        recruiter_note(),
        encoding="utf-8",
    )

    package = ROOT / "TrendScale_v5_Deployable_Send_Package.zip"
    package_items = [
        ROOT / "TrendScale_JCKED_Production_Brief_FINAL.docx",
        ROOT / "TrendScale_Puravita_Production_Brief_FINAL.docx",
        ROOT / "TrendScale_v5_Recruiter_Revision_Note.md",
    ]
    with zipfile.ZipFile(package, "w", zipfile.ZIP_DEFLATED) as archive:
        for item in package_items:
            archive.write(item, item.name)


def main() -> None:
    briefs = [JCKED, PURAVITA]
    paths = [build_docx(brief) for brief in briefs]
    write_outputs(paths, briefs)
    for path in paths:
        print(path)
    print(ROOT / "TrendScale_Production_Briefs_v5_Deployable_Text_Extract.md")
    print(ROOT / "TrendScale_v5_Recruiter_Revision_Note.md")
    print(ROOT / "TrendScale_v5_Deployable_Script_Quality_Snapshot.md")
    print(ROOT / "TrendScale_v5_Script_RedTeam_Receipt.md")
    print(ROOT / "TrendScale_v5_Deployable_Send_Package.zip")


if __name__ == "__main__":
    main()
