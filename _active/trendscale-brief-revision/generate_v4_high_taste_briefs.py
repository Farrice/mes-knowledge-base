from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from generate_v3_production_briefs import ROOT, TEMPLATE, docx_to_markdown, set_cell, set_para


def add_row(table, values: list[str]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def replace_table_rows(table, rows: list[list[str]]) -> None:
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for row in rows:
        add_row(table, row)


def build_docx(brief: dict) -> Path:
    doc = Document(TEMPLATE)

    paragraph_values = {
        1: "Client-facing production brief. v4 high-taste elevation pass for founder review and editor handoff.",
        3: f"Strategist: {brief['strategist']}",
        4: f"Client: {brief['client']}",
        5: f"PDP: {brief['pdp']}",
        6: f"Media buyer notes: {brief['media_buyer_notes']}",
        7: f"Concept type: {brief['concept_type']}",
        8: f"Iterated from: {brief['iterated_from']}",
        9: f"Avatar: {brief['avatar']}",
        10: f"Awareness level: {brief['awareness']}",
        11: f"Hypothesis: {brief['hypothesis']}",
        13: f"Concept inspo: {brief['concept_inspo']}",
        14: f"Cast type: {brief['cast_type']}",
        15: f"AI avatar look: {brief['ai_avatar']}",
        16: f"Pacing: {brief['pacing']}",
        17: f"Main emotion: {brief['main_emotion']}",
        18: f"Voiceover: {brief['voiceover']}",
        19: f"Music: {brief['music']}",
        20: f"Caption style: {brief['caption_style']}",
        21: f"Overall direction: {brief['overall_direction']}",
        22: f"Editing notes: {brief['editing_notes']}",
        24: f"Format: {brief['format']}",
        26: f"Note: {brief['note']}",
    }

    for index, text in paragraph_values.items():
        set_para(doc.paragraphs[index], text)

    replace_table_rows(doc.tables[0], brief["script_rows"])

    out_path = ROOT / brief["filename"]
    doc.save(out_path)
    return out_path


JCKED = {
    "filename": "TrendScale_JCKED_Production_Brief_v4.docx",
    "strategist": "Farrice Cain",
    "client": "JCKED - Liquid L-Carnitine 4000mg",
    "pdp": "https://jcked.com/products/liquid-l-carnitine-4000mg-of1",
    "media_buyer_notes": (
        "Primary test: prior-bottle skepticism into serving-size contrast. Make "
        "4,000mg the protected idea. Keep competitor comparison generic unless "
        "specific labels are verified."
    ),
    "concept_type": "Iteration",
    "iterated_from": "The Locked Vault control concept.",
    "avatar": (
        "Male, 35 to 50. He has a cabinet graveyard of half-used supplements and "
        "does not trust transformation language. He will stop for a clean label "
        "contradiction: maybe the ingredient was never the real test."
    ),
    "awareness": "Problem aware to Solution aware",
    "hypothesis": (
        "The vault metaphor works when it reopens an old purchase. Start with the "
        "bottle he already dismissed, turn on serving size, then make the label "
        "comparison feel unavoidable."
    ),
    "concept_inspo": (
        "Half-used bottle, bathroom cabinet, locked vault, amber key, mitochondrial "
        "gate, serving-size cards, real product on counter."
    ),
    "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
    "ai_avatar": (
        "Male, 38 to 45, understated, clean but lived-in. Neutral kitchen or study. "
        "Simple dark overshirt or crewneck. Calm strategist energy, no doctor coat."
    ),
    "pacing": "Cold object open, quick dose contradiction, measured mechanism, slow product hold.",
    "main_emotion": "Vindication. The viewer should feel the old bottle may have been a weak test, not a final answer.",
    "voiceover": "Low, plain, and forensic. It should feel like a label being reopened, not a pitch being performed.",
    "music": "Sparse low pulse at minus 18dB. Drop under the final label line.",
    "caption_style": "Manual burned-in captions. Six words per frame maximum. Amber only for key, gate, and dose-card moments.",
    "overall_direction": (
        "Open with the object. Let the vault metaphor earn the mechanism. Land on "
        "the label. The whole ad should make one sentence feel obvious: compare "
        "the serving before judging the ingredient."
    ),
    "editing_notes": (
        "Use cabinet, label, dose card, vault, and mitochondria visuals. Avoid body "
        "photos, scale imagery, before-after comparisons, doctor visuals, panic "
        "urgency, and template effects. Keep the bottle real."
    ),
    "format": "Video, 9:16, 30 seconds. Optional 6-card static carousel using the serving-size contrast.",
    "note": (
        "Structure-function language only. Final edit must match the PDP and label. "
        "Avoid hard fat-loss promises, disease claims, personal-attribute callouts, "
        "and unverified competitor claims. Source anchors: NIH ODS Carnitine "
        "Health Professional Fact Sheet, https://ods.od.nih.gov/factsheets/"
        "Carnitine-HealthProfessional/; Linus Pauling Institute L-Carnitine, "
        "https://lpi.oregonstate.edu/mic/dietary-factors/L-carnitine."
    ),
    "script_rows": [
        [
            "Half-used L-carnitine bottle in a bathroom cabinet. Thumb turns the label toward camera.",
            "Hook 1",
            "He tried L-carnitine once. The ingredient got blamed. The serving size never got questioned.",
            "QUESTION THE SERVING",
        ],
        [
            "Locked amber vault inside a stylized cell. A key stops just short of the gate.",
            "Hook 2",
            "Stored fuel can sit right there and still be locked away.",
            "THE DOOR STAYS SHUT",
        ],
        [
            "Two clean cards: 500mg on the left, JCKED 4,000mg on the right. No competitor logo.",
            "Hook 3",
            "A 500mg label and a 4,000mg label deserve different judgments.",
            "500MG / 4,000MG",
        ],
        [
            "Cinematic mitochondria gate. L-carnitine tag moves with a long-chain fatty acid.",
            "Mechanism",
            "Long-chain fatty acids need help reaching mitochondria. L-carnitine is part of that transport system.",
            "TRANSPORT SYSTEM",
        ],
        [
            "Dose card returns. Prior bottle fades back. JCKED 4,000mg stays sharp.",
            "Dose turn",
            "That is why the serving matters. A familiar ingredient can still be a weak test.",
            "SERVING SIZE MATTERS",
        ],
        [
            "Real JCKED bottle on a counter. Amber key line points to the 4,000mg claim.",
            "Product reveal",
            "JCKED puts 4,000mg liquid L-carnitine on the label. That is the number to compare.",
            "4,000MG LIQUID L-CARNITINE",
        ],
        [
            "Bottle hold. Quick label comparison card. CTA appears after the voice lands.",
            "CTA",
            "Read the label. Compare the serving. Then decide whether the first test was real.",
            "COMPARE THE SERVING",
        ],
    ],
}


PURAVITA = {
    "filename": "TrendScale_Puravita_Production_Brief_v4.docx",
    "strategist": "Farrice Cain",
    "client": "Puravita - 12-Form Magnesium Complex",
    "pdp": "https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex",
    "media_buyer_notes": (
        "Primary test: hidden-signal metaphor into form-list curiosity. Lead with "
        "phone or lab-report tension, then move to 300+ enzyme systems, serum "
        "limits, form differences, and the 12-form PDP story."
    ),
    "concept_type": "Iteration",
    "iterated_from": "The Battery You Can't See control concept.",
    "avatar": (
        "Men 38 to 62, with a 45 to 55 sweet spot. He has routines, labs, and "
        "enough life pressure to distrust another sleep-magnesium ad. He needs a "
        "clean way to understand the signal before he considers the bottle."
    ),
    "awareness": "Unaware to Problem aware",
    "hypothesis": (
        "The battery concept works when the first beat feels like a missing "
        "dashboard, not a symptom list. Once the signal is open, the serum window "
        "and 12-form label give Puravita a reason to be inspected."
    ),
    "concept_inspo": (
        "Phone at 5 percent, normal lab panel, untouched coffee, soft body-map glow, "
        "serum window card, form list, quiet product reveal."
    ),
    "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
    "ai_avatar": (
        "Male, 45 to 50, understated, slight gray at temples, simple sweater or "
        "overshirt, warm kitchen light. Calm and observant."
    ),
    "pacing": "Quiet first beat, sharp signal turn, clean proof card, slow label reveal.",
    "main_emotion": "Recognition without diagnosis. The viewer should feel the ad named a gap without naming him.",
    "voiceover": "Quiet, restrained, and observational. It should sound like someone noticing the missing instrument on the dashboard.",
    "music": "Sparse piano or low ambient bed at minus 18dB. Silence under the final label line.",
    "caption_style": "Manual burned-in captions. Six words per frame maximum. Sage for serum and form callouts.",
    "overall_direction": (
        "Make the first half an invisible-signal story and the second half a label "
        "inspection story. Keep the battery metaphor, but let the product earn the "
        "click through the form list."
    ),
    "editing_notes": (
        "Use hand-only phone b-roll and clean document-style lab visuals. Use a "
        "soft body-map glow, not literal anatomy. Avoid exhausted-person stock, "
        "doctor coats, disease language, diagnosis, urgency text, and influencer "
        "quote cards unless cleared by the client."
    ),
    "format": "Video, 9:16, 35 to 45 seconds. Optional 4:5 static thumbstop using the phone or lab-report hook.",
    "note": (
        "Use label/PDP-approved structure-function claims only. Correct proof base: "
        "magnesium supports more than 300 enzyme systems; less than 1% of total "
        "body magnesium is in serum; supplement forms can differ in absorption. "
        "Avoid depletion certainty, sleep cures, fatigue claims, and unapproved "
        "named-influencer usage. Source anchor: NIH ODS Magnesium Health "
        "Professional Fact Sheet, https://ods.od.nih.gov/factsheets/"
        "Magnesium-HealthProfessional/."
    ),
    "script_rows": [
        [
            "Phone at 5 percent on a nightstand. Dawn light. Hand only, no face.",
            "Hook 1",
            "A phone shows 5 percent. The body rarely gets a battery icon.",
            "NO BODY BATTERY ICON",
        ],
        [
            "Clean lab report slides into frame. The word normal is visible. Coffee sits untouched.",
            "Hook 2",
            "The report says normal. The morning is less polite.",
            "NORMAL ON PAPER",
        ],
        [
            "Common sleep-ad visuals flicker off. A soft body-map glow replaces them.",
            "Hook 3",
            "Sleep gets the headline. The signal comes first.",
            "BEFORE SLEEP: SIGNAL",
        ],
        [
            "Body map lights up nerves, muscle, and energy icons. Keep it editorial, not medical.",
            "Mechanism",
            "Magnesium supports more than 300 enzyme systems. Energy metabolism. Nerves. Muscle.",
            "300+ ENZYME SYSTEMS",
        ],
        [
            "Bloodstream label shows less than 1%. Camera pulls back to bone, muscle, and soft tissue.",
            "Proof",
            "Less than 1% of total body magnesium is in serum. That makes bloodwork a narrow window.",
            "<1% IN SERUM",
        ],
        [
            "Formula board. Twelve form cards appear one by one, clean and readable.",
            "Form turn",
            "After the signal, look at the form. Magnesium supplements do not absorb the same way.",
            "FORM CHANGES THE STORY",
        ],
        [
            "Real Puravita bottle. Twelve form cards settle around the label.",
            "Product reveal",
            "Puravita brings 12 magnesium forms into one formula. That makes the label worth reading.",
            "12-FORM MAGNESIUM COMPLEX",
        ],
        [
            "Bottle and supplement-facts panel. CTA appears after the final sentence.",
            "CTA",
            "Read the form list. Check the amount. Choose with the label open.",
            "READ THE FORM LIST",
        ],
    ],
}


def sentence_lengths(text: str) -> list[int]:
    parts = [p.strip() for p in re.split(r"[.!?]+", text) if p.strip()]
    return [len(re.findall(r"\b[\w'<%+]+\b", p)) for p in parts]


def quality_snapshot(briefs: list[dict]) -> str:
    lines = ["# TrendScale v4 High-Taste Quality Snapshot", ""]
    for brief in briefs:
        script_texts = [row[2] for row in brief["script_rows"]]
        lengths = [length for text in script_texts for length in sentence_lengths(text)]
        avg = round(sum(lengths) / len(lengths), 1)
        max_len = max(lengths)
        row_word_counts = [len(re.findall(r"\b[\w'<%+]+\b", text)) for text in script_texts]
        lines.append(f"## {brief['client']}")
        lines.append(f"- Script rows: {len(script_texts)}")
        lines.append(f"- Average sentence length: {avg} words")
        lines.append(f"- Longest sentence: {max_len} words")
        lines.append(f"- Longest script row: {max(row_word_counts)} words")
        lines.append("")
    lines.append("## High-Taste Moves")
    lines.append("- Restored the original cinematic object tension: old bottle, locked door, low battery, normal report.")
    lines.append("- Replaced clipped v3 lines with visual, voice-ready script beats.")
    lines.append("- Kept PDP links live inside both briefs.")
    lines.append("- Preserved source-backed claims and launch-safe wording.")
    lines.append("- Made on-screen text more ownable and less generic.")
    lines.append("")
    return "\n".join(lines)


def write_outputs(paths: list[Path], briefs: list[dict]) -> None:
    extract = ["# TrendScale Production Briefs v4 High-Taste Text Extract", ""]
    for path in paths:
        extract.append(docx_to_markdown(path))
        extract.append("")
    (ROOT / "TrendScale_Production_Briefs_v4_High_Taste_Text_Extract.md").write_text(
        "\n".join(extract),
        encoding="utf-8",
    )

    (ROOT / "TrendScale_v4_High_Taste_Quality_Snapshot.md").write_text(
        quality_snapshot(briefs),
        encoding="utf-8",
    )

    send_note = """Hi [Recruiter Name],

Thanks again for the feedback. I made one more pass after adding the PDP links and tightened both briefs at the script level.

The concepts are still the same: JCKED is The Locked Vault, and Puravita is The Battery You Can't See. The difference is that the hooks, body sequence, on-screen text, and production notes now read more like production-ready paid creative rather than a strategy draft.

Attached:
1. TrendScale_JCKED_Production_Brief_v4.docx
2. TrendScale_Puravita_Production_Brief_v4.docx

Best,
Farrice
"""
    (ROOT / "TrendScale_v4_Recruiter_Revision_Note.md").write_text(send_note, encoding="utf-8")


def main() -> None:
    briefs = [JCKED, PURAVITA]
    paths = [build_docx(brief) for brief in briefs]
    write_outputs(paths, briefs)
    for path in paths:
        print(path)
    print(ROOT / "TrendScale_Production_Briefs_v4_High_Taste_Text_Extract.md")
    print(ROOT / "TrendScale_v4_Recruiter_Revision_Note.md")
    print(ROOT / "TrendScale_v4_High_Taste_Quality_Snapshot.md")


if __name__ == "__main__":
    main()
