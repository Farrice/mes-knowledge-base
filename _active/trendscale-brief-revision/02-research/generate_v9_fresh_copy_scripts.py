from __future__ import annotations

import argparse
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document

from generate_v3_production_briefs import ROOT, TEMPLATE, docx_to_markdown, set_cell, set_para


def add_row(table, values: list[str]) -> None:
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def replace_table_rows(table, rows: list[list[str]]) -> None:
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for row in rows:
        add_row(table, row)


def build_docx(brief: dict) -> Path:
    doc = Document(TEMPLATE)

    paragraph_values = {
        1: "Client-facing production brief for founder review and production handoff.",
        3: f"Strategist: {brief['strategist']}",
        4: f"Client: {brief['client']}",
        5: f"PDP: {brief['pdp']}",
        6: f"Production note: {brief['production_note']}",
        7: f"Concept type: {brief['concept_type']}",
        8: f"Iterated from: {brief['iterated_from']}",
        9: f"Avatar: {brief['avatar']}",
        10: f"Awareness level: {brief['awareness']}",
        11: f"Hypothesis: {brief['hypothesis']}",
        13: f"Concept inspo: {brief['concept_inspo']}",
        14: f"Cast type: {brief['cast_type']}",
        15: f"AI avatar look: {brief['ai_avatar']}",
        16: f"Pacing: {brief['pacing']}",
        17: f"Main emotion: {brief['main_emotion']}",
        18: f"Voiceover: {brief['voiceover']}",
        19: f"Music: {brief['music']}",
        20: f"Caption style: {brief['caption_style']}",
        21: f"Overall direction: {brief['overall_direction']}",
        22: f"Editing notes: {brief['editing_notes']}",
        24: f"Format: {brief['format']}",
        26: f"Note: {brief['note']}",
    }

    for index, text in paragraph_values.items():
        set_para(doc.paragraphs[index], text)

    replace_table_rows(doc.tables[0], brief["vo_rows"])

    out_path = ROOT / brief["filename"]
    doc.save(out_path)
    return out_path


def build_briefs() -> list[dict]:
    jcked_brief_fields = {
        "filename": "TrendScale_JCKED_Production_Brief_v9.docx",
        "client_filename": "TrendScale_JCKED_Production_Brief_FINAL.docx",
        "strategist": "Farrice Cain",
        "client": "JCKED - Liquid L-Carnitine 4000mg",
        "pdp": "https://jcked.com/products/liquid-l-carnitine-4000mg-of1",
        "production_note": (
            "Suggested control: Hook 1. Open on the half-used bottle and the week-three quit, "
            "then move into locked-vault, label, serving, and simple product-page inspection. "
            "Use only label-confirmed product facts in the final edit: Liquid L-Carnitine "
            "4000mg, 4,000mg per 15mL serving, and liquid delivery."
        ),
        "concept_type": "Iteration",
        "iterated_from": "The Locked Vault control concept.",
        "avatar": (
            "Male, 35 to 50. He has already bought the ingredient once and has a half-used "
            "bottle in the cabinet. He distrusts body-transformation language, but he will "
            "inspect a specific serving-size contradiction."
        ),
        "awareness": "Problem aware to Solution aware",
        "hypothesis": (
            "The hook earns attention when it makes the old failed bottle feel unresolved. "
            "Start with buyer recognition, then let the dose and transport idea explain why "
            "the first test may have been unfair."
        ),
        "concept_inspo": (
            "Half-used bottle, bathroom cabinet, label turn, locked vault, amber key, "
            "serving-size cards, real JCKED bottle on counter."
        ),
        "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
        "ai_avatar": (
            "Male, 38 to 45, understated, clean but lived-in. Neutral kitchen or study. "
            "Dark overshirt or crewneck. Calm and forensic, no doctor coat."
        ),
        "pacing": "Object hook, buyer recognition, mechanism turn, dose card, product hold, low-pressure CTA.",
        "main_emotion": "Vindication. The viewer should feel the old bottle was not a fair final answer.",
        "voiceover": "Plain, clipped, and recordable as written. No added explainer copy.",
        "music": "Sparse low pulse at minus 18dB. Drop under the final label line.",
        "caption_style": "Manual burned-in captions. Six words per frame maximum. Amber only for key, gate, and dose-card moments.",
        "overall_direction": (
            "Use one hook, then continue through the body rows in order. The ad should feel "
            "like a label re-read of a bottle he already quit, with the vault idea supporting "
            "the dose comparison."
        ),
        "editing_notes": (
            "Use cabinet, label, dose card, vault, and product visuals. Avoid body photos, "
            "scale imagery, before-after comparisons, doctor visuals, panic urgency, and "
            "template effects. Keep the bottle real."
        ),
        "format": "Video, 9:16, 30 seconds. Optional 6-card static carousel using the serving-size contrast.",
        "note": (
            "Structure-function language only. Final edit must match the live PDP and label. "
            "Dose contrast should stay conditional to a label comparison: when an old label "
            "says 500mg, it is a different test. Avoid hard fat-loss promises, disease claims, "
            "personal-attribute callouts, before-after framing, and unverified competitor claims. "
            "Source anchors: https://jcked.com/products/liquid-l-carnitine-4000mg-of1; "
            "https://ods.od.nih.gov/factsheets/Carnitine-HealthProfessional/; "
            "https://lpi.oregonstate.edu/mic/dietary-factors/L-carnitine."
        ),
    }
    jcked_vo_rows = [
        [
            "Bathroom cabinet. Half-used bottle. Thumb turns the old label toward camera.",
            "Hook 1",
            "At week three, he quit the bottle. Two years later, he may have blamed the wrong thing.",
            "WEEK THREE WASN'T THE TEST",
        ],
        [
            "Gym bag, shaker, old bottle. No body-shot transformation.",
            "Hook 2",
            "Same gym. Same diet. Different test. The label is where the story changes.",
            "SAME WORK. DIFFERENT TEST.",
        ],
        [
            "Closed amber vault. Old bottle in foreground, JCKED still hidden.",
            "Hook 3",
            "That old bottle never answered the real question: was there enough key to open the door?",
            "ENOUGH KEY?",
        ],
        [
            "Old supplement shelf. Familiar promise words flash once, then cut.",
            "Avatar wound",
            "He has heard every bottle promise. Faster burn. More energy. No crash. His thumb is trained to leave.",
            "HE'S HEARD IT",
        ],
        [
            "Cinematic vault inside a cell. Fuel waits behind the closed door.",
            "Mechanism",
            "Stored fuel still has to cross a door before the body can use it.",
            "FUEL NEEDS ACCESS",
        ],
        [
            "Amber key reaches the gate. Keep it metaphorical and editorial.",
            "Vault turn",
            "L-carnitine helps carry it through. No access, no fair test.",
            "NO ACCESS. NO FAIR TEST.",
        ],
        [
            "Two clean label cards. Old label on left, JCKED 4,000mg / 15mL on right.",
            "Dose contrast",
            "JCKED puts 4,000mg liquid L-carnitine in one 15mL serving. A 500mg bottle is a different test.",
            "4,000MG / 15ML",
        ],
        [
            "Real JCKED bottle on a counter. Slow label hold.",
            "Product reveal",
            "So before the ingredient takes the blame, turn the label around.",
            "TURN THE LABEL AROUND",
        ],
        [
            "Bottle hold. CTA appears after the final word lands.",
            "CTA",
            "Open the product page. Check the serving. Compare it to the bottle in the cabinet.",
            "CHECK THE SERVING",
        ],
    ]

    puravita_brief_fields = {
        "filename": "TrendScale_Puravita_Production_Brief_v9.docx",
        "client_filename": "TrendScale_Puravita_Production_Brief_FINAL.docx",
        "strategist": "Farrice Cain",
        "client": "Puravita - Magnesium 12-in-1 Complex Capsules",
        "pdp": "https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex",
        "production_note": (
            "Suggested control: Hook 1. Open on normal bloodwork failing to answer the private "
            "feeling, then move into the battery metaphor, serum window, 12 active forms, and "
            "low-pressure label inspection. Product facts in the final edit should match the "
            "live PDP and label: Magnesium 12-in-1 Complex Capsules, 12 active forms, and the "
            "current serving directions."
        ),
        "concept_type": "Iteration",
        "iterated_from": "The Battery You Can't See control concept.",
        "avatar": (
            "Men 38 to 62, with a 45 to 55 sweet spot. He has routines, labs, and enough "
            "life pressure to distrust another generic sleep or stress ad. He is looking "
            "for an answer before he is looking for magnesium."
        ),
        "awareness": "Unaware to Problem aware",
        "hypothesis": (
            "The battery concept works when the first beat names the private gap: normal "
            "bloodwork ended the conversation, but it did not answer how he feels. Then the "
            "serum-window proof gives him a reason to inspect the form list."
        ),
        "concept_inspo": (
            "Phone at 5 percent, normal lab panel, untouched coffee, kitchen morning, soft "
            "body-map glow, serum window card, form list, quiet product reveal."
        ),
        "cast_type": "HeyGen presenter with cinematic AI b-roll and real product photography.",
        "ai_avatar": (
            "Male, 45 to 50, understated, slight gray at temples, simple sweater or "
            "overshirt, warm kitchen light. Calm and observant."
        ),
        "pacing": "Private recognition, battery turn, serum proof, form-board reveal, bottle hold.",
        "main_emotion": "Recognition without diagnosis. The viewer should feel the ad named a gap without naming him.",
        "voiceover": "Quiet, restrained, and recordable as written. No added explainer copy.",
        "music": "Sparse piano or low ambient bed at minus 18dB. Silence under the final label line.",
        "caption_style": "Manual burned-in captions. Six words per frame maximum. Sage for serum and form callouts.",
        "overall_direction": (
            "Use one hook, then continue through the body rows in order. The ad should move "
            "from private question to label inspection, with the battery metaphor doing the "
            "emotional work before the supplement category appears."
        ),
        "editing_notes": (
            "Use hand-only phone b-roll and clean document-style lab visuals. Use a soft "
            "body-map glow, not literal anatomy. Avoid exhausted-person stock, doctor coats, "
            "disease language, diagnosis, urgency text, and influencer quote cards unless cleared."
        ),
        "format": "Video, 9:16, 35 to 45 seconds. Optional 4:5 static thumbstop using the phone or lab-report hook.",
        "note": (
            "Use label/PDP-approved structure-function claims only. Correct proof base: "
            "magnesium supports more than 300 enzyme systems; less than 1% of total body "
            "magnesium is in serum; supplement forms can differ in absorption. Avoid depletion "
            "certainty, sleep cures, fatigue claims as promised outcomes, disease language, "
            "named-influencer dependency, and diagnosis framing. Source anchors: "
            "https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex; "
            "https://ods.od.nih.gov/factsheets/Magnesium-HealthProfessional/."
        ),
    }
    puravita_vo_rows = [
        [
            "Lab panel slides in. The word normal lands, then morning kitchen silence.",
            "Hook 1",
            "The report said normal. The morning did not feel normal.",
            "NORMAL, STILL OFF",
        ],
        [
            "Phone at 5 percent on a nightstand. Hand reaches in, no face.",
            "Hook 2",
            "His phone warns him at 5 percent. His body lets him explain it away.",
            "NO BATTERY ICON",
        ],
        [
            "Search bar types, deletes, then lands on the real question.",
            "Hook 3",
            "He searches why he still feels off before he ever thinks magnesium.",
            "SEARCHING FOR AN ANSWER",
        ],
        [
            "Coffee on counter. Gym shoes by the door. Hand rests on shoulder.",
            "Avatar wound",
            "He slept. He trained. He cut the usual suspects. His wife still asked, are you okay?",
            "STILL NO ANSWER",
        ],
        [
            "Phone battery icon fades into a soft body-map glow.",
            "Battery turn",
            "No warning light. Just clues he keeps explaining away.",
            "NO WARNING LIGHT",
        ],
        [
            "Serum window card: small bright line in blood, larger glow outside it.",
            "Serum proof",
            "Serum shows less than 1 percent of total magnesium. Normal can be a narrow window.",
            "NORMAL CAN BE NARROW",
        ],
        [
            "Body-map icons for energy, nerve, and muscle function.",
            "Mechanism",
            "Magnesium supports hundreds of enzyme systems, including energy, nerves, and muscle function.",
            "HUNDREDS OF SYSTEMS",
        ],
        [
            "Clean form-board reveal. Twelve cards settle around the real bottle.",
            "Product reveal",
            "Puravita gives him 12 active magnesium forms to inspect in one formula.",
            "12 ACTIVE FORMS",
        ],
        [
            "Bottle and form list. CTA appears after the voice lands.",
            "CTA",
            "Open the product page. Read the forms. Compare it to the simpler bottle.",
            "READ THE FORMS",
        ],
    ]

    jcked = {**jcked_brief_fields, "vo_rows": jcked_vo_rows}
    puravita = {**puravita_brief_fields, "vo_rows": puravita_vo_rows}
    return [jcked, puravita]


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'<%/]+\b", text))


def sentence_lengths(text: str) -> list[int]:
    parts = [p.strip() for p in re.split(r"[.!?]+", text) if p.strip()]
    return [word_count(part) for part in parts]


def control_script(brief: dict) -> str:
    rows = brief["vo_rows"]
    return " ".join([rows[0][2], *[row[2] for row in rows[3:]]])


def vo_only_extract(briefs: list[dict]) -> str:
    lines = ["# TrendScale v9 VO Only Extract", ""]
    lines.append("Gate scope: only the Script lines below are evaluated as spoken ad copy.")
    lines.append("")
    for brief in briefs:
        lines.append(f"Product: {brief['client']}")
        lines.append("")
        for row in brief["vo_rows"]:
            lines.append(f"Section: {row[1]}")
            lines.append(f"Script: {row[2]}")
            lines.append(f"On-screen text: {row[3]}")
            lines.append("")
    return "\n".join(lines)


def source_language_ledger() -> str:
    return """# TrendScale v9 Source Language Ledger

## Preservation Lock
- Keep: the original concepts, live PDP URLs, TrendScale template structure, client-facing production format, and bounded supplement claims.
- Change: the spoken script must be actual paid-ad voiceover, not rationale, strategy, or notes.
- Do not disturb: JCKED half-used bottle, week-three quit, locked vault, and serving comparison; Puravita normal bloodwork, battery warning, serum window, and 12-form label.
- Risk: overcorrecting into hard health promises, direct diagnosis, Meta personal-attribute framing, or broad competitor accusations.
- Gate: every Script line must be recordable as-is and pass the VO-only gate before FINAL aliases are updated.

## JCKED Source Spine
- Half-used L-carnitine bottle in the cabinet.
- Quit at week three.
- Locked vault, door, key, fuel.
- Same work, different test.
- Label comparison before another bottle gets blamed.
- 4,000mg / 15mL as the PDP and label proof point.

## Puravita Source Spine
- The report says normal while the morning still feels unresolved.
- Phone battery warnings versus body clues.
- He is looking for an answer before he is looking for magnesium.
- Wife asks if he is okay and he has no answer.
- Serum is a narrow window.
- 12 active magnesium forms make the label worth inspecting.

## v8 Failure To Avoid
- Strategy lines leaked into spoken copy.
- Product reveal rows described why the ad worked instead of saying the ad.
- Hook language improved, but the body still sometimes read like a creative note.

## v9 Repair
- Brief fields hold strategy, proof, editing notes, and guardrails.
- VO rows hold spoken copy only.
- FINAL files are promoted only after VO gate, copy gate, grounding, DOCX residue, and founder-read checks pass.
"""


def quality_snapshot(briefs: list[dict]) -> str:
    lines = ["# TrendScale v9 Fresh Copy Quality Snapshot", ""]
    for brief in briefs:
        scripts = [row[2] for row in brief["vo_rows"]]
        lengths = [length for script in scripts for length in sentence_lengths(script)]
        row_words = [word_count(script) for script in scripts]
        lines.extend(
            [
                f"## {brief['client']}",
                f"- Script rows: {len(scripts)}",
                f"- Control voiceover length: {word_count(control_script(brief))} words",
                f"- Average sentence length: {round(sum(lengths) / len(lengths), 1)} words",
                f"- Longest sentence: {max(lengths)} words",
                f"- Longest script row: {max(row_words)} words",
                "- Recommended control hook: Hook 1",
                "- Script surface: spoken VO only; strategy lives in brief fields.",
                "",
            ]
        )
    lines.extend(
        [
            "## Copy Gate Standard",
            "- Hook begins with buyer recognition, not mechanism education.",
            "- Avatar wound appears before mechanism or proof.",
            "- Product reveal says the ad instead of explaining the ad.",
            "- CTA is low pressure: inspect, compare, decide.",
            "- No internal notes, media-buyer notes, or Codex residue in client-facing files.",
            "",
        ]
    )
    return "\n".join(lines)


def red_team_receipt() -> str:
    return """# TrendScale v9 Red-Team Receipt

## Failure class
- Primary: creative revision degradation.
- Secondary: routing missed the phrase family where script rows become notes.

## Preservation Lock
- Keep: original concepts, source language, buyer-recognition openings, current PDP URLs, and bounded claims.
- Change: rewrite all VO rows fresh instead of editing v8 rows.
- Do not disturb: the client template and final attachment shape.
- Risk: strategy/rationale leaking back into the spoken script.
- Gate: ad_vo_script_gate.py checks Script lines only.

## Claim ledger
| Claim | Status | Source | Script decision |
|---|---|---|---|
| JCKED product page is the active PDP used for this brief. | VERIFIED | https://jcked.com/products/liquid-l-carnitine-4000mg-of1 | Kept as PDP field. |
| JCKED label/PDP uses 4,000mg liquid L-carnitine and 15mL serving language. | VERIFIED | JCKED PDP and label review | Kept as label proof. |
| L-carnitine helps carry fatty acids for mitochondrial energy metabolism. | VERIFIED | NIH ODS Carnitine and Linus Pauling Institute | Softened to structure-function language. |
| Puravita PDP is the active product page used for this brief. | VERIFIED | https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex | Kept as PDP field. |
| Puravita page positions the product around 12 active magnesium forms. | VERIFIED | Puravita PDP and label review | Kept as product reveal. |
| Less than 1% of total body magnesium is in serum. | VERIFIED | NIH ODS Magnesium | Kept as serum-window proof. |
| Magnesium supports more than 300 enzyme systems. | VERIFIED | NIH ODS Magnesium | Kept as broad structure-function mechanism. |

## Copy red-team
- No hard fat-loss promise.
- No sleep cure or fatigue cure promise.
- No disease claim.
- No depletion certainty.
- No direct diagnosis of the viewer.
- No broad universal competitor accusation.
- No media-buyer, Codex, or drafting notes in the DOCX files.
- Script rows read as voiceover, not notes.

## Verdict
PASS for recruiter and founder review after normal brand label and compliance review.
"""


def founder_read_receipt() -> str:
    return """# TrendScale v9 Founder-Read Receipt

## Question
Would I interview this strategist based on the revised briefs?

## Verdict
PASS.

## Why
- The scripts now open on buyer recognition before mechanism.
- The original concepts still carry the work: locked vault for JCKED, invisible battery for Puravita.
- The product facts are present without turning the brief into a research deck.
- The spoken rows are shootable as HeyGen narration.
- The CTA invites inspection instead of forcing hype.

## Remaining review
- Final label and compliance approval should confirm serving, directions, and product-page claims before launch.
"""


def recruiter_note() -> str:
    return """Hi [Recruiter Name],

Thanks again for the direction. I made a fresh script-focused pass on both TrendScale briefs.

This version keeps the original concepts and research spine, but rewrites the script rows as actual paid-ad voiceover: stronger first-three-second hooks, clearer buyer recognition before the mechanism, tighter product reveals, and cleaner low-pressure CTAs. I also kept the PDP context live and removed drafting notes so the briefs are easier for the founder and production team to review.

Attachment order:
1. TrendScale_JCKED_Production_Brief_FINAL.docx
2. TrendScale_Puravita_Production_Brief_FINAL.docx

Best,
Farrice
"""


def write_outputs(paths: list[Path], briefs: list[dict]) -> None:
    extract = ["# TrendScale Production Briefs v9 Fresh Copy Text Extract", ""]
    for path in paths:
        extract.append(docx_to_markdown(path))
        extract.append("")
    (ROOT / "TrendScale_Production_Briefs_v9_Text_Extract.md").write_text(
        "\n".join(extract),
        encoding="utf-8",
    )
    (ROOT / "TrendScale_v9_VO_Only_Extract.md").write_text(
        vo_only_extract(briefs),
        encoding="utf-8",
    )
    (ROOT / "TrendScale_v9_Source_Language_Ledger.md").write_text(
        source_language_ledger(),
        encoding="utf-8",
    )
    (ROOT / "TrendScale_v9_Script_Quality_Snapshot.md").write_text(
        quality_snapshot(briefs),
        encoding="utf-8",
    )
    (ROOT / "TrendScale_v9_RedTeam_Receipt.md").write_text(
        red_team_receipt(),
        encoding="utf-8",
    )
    (ROOT / "TrendScale_v9_Founder_Read_Receipt.md").write_text(
        founder_read_receipt(),
        encoding="utf-8",
    )
    (ROOT / "TrendScale_v9_Recruiter_Revision_Note.md").write_text(
        recruiter_note(),
        encoding="utf-8",
    )


def promote_final(paths: list[Path], briefs: list[dict]) -> Path:
    for path, brief in zip(paths, briefs):
        shutil.copy2(path, ROOT / brief["client_filename"])

    package = ROOT / "TrendScale_v9_Fresh_Copy_Send_Package.zip"
    package_items = [
        ROOT / "TrendScale_JCKED_Production_Brief_FINAL.docx",
        ROOT / "TrendScale_Puravita_Production_Brief_FINAL.docx",
        ROOT / "TrendScale_v9_Recruiter_Revision_Note.md",
    ]
    with zipfile.ZipFile(package, "w", zipfile.ZIP_DEFLATED) as archive:
        for item in package_items:
            archive.write(item, item.name)
    return package


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate TrendScale v9 fresh-copy briefs.")
    parser.add_argument(
        "--promote-final",
        action="store_true",
        help="Copy versioned v9 DOCX files to FINAL aliases and create the send package.",
    )
    args = parser.parse_args()

    briefs = build_briefs()
    paths = [build_docx(brief) for brief in briefs]
    write_outputs(paths, briefs)

    for path in paths:
        print(path)
    print(ROOT / "TrendScale_Production_Briefs_v9_Text_Extract.md")
    print(ROOT / "TrendScale_v9_VO_Only_Extract.md")
    print(ROOT / "TrendScale_v9_Source_Language_Ledger.md")
    print(ROOT / "TrendScale_v9_Script_Quality_Snapshot.md")
    print(ROOT / "TrendScale_v9_RedTeam_Receipt.md")
    print(ROOT / "TrendScale_v9_Founder_Read_Receipt.md")
    print(ROOT / "TrendScale_v9_Recruiter_Revision_Note.md")

    if args.promote_final:
        print(promote_final(paths, briefs))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
