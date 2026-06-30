#!/usr/bin/env python3
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path("/Users/farricecain/Downloads/TrendScale_Master_Brief_Template (1).docx")


BRIEFS = {
    "jcked": {
        "out": ROOT / "TrendScale_JCKED_Production_Brief_v10.docx",
        "product": "JCKED",
        "fields": {
            4: "Strategist: Farrice Cain",
            5: "Client: JCKED Liquid L-Carnitine 4000mg",
            6: "PDP: https://jcked.com/products/liquid-l-carnitine-4000mg-of1",
            7: "Media buyer notes: Quiet editorial cold-traffic test. Avoid urgency, before-after visuals, hard fat-loss promises, clinical-dose language, and competitor dose claims unless the brand clears them.",
            8: "Concept type: Iteration",
            9: "Iterated from: The Locked Vault",
            10: "Avatar: Male 32 to 45 who already tried L-carnitine, left a half-used bottle in the cabinet, still trains and tracks, and wants a fairer explanation than discipline.",
            11: "Awareness level: Problem to Solution",
            12: "Hypothesis: If the ad starts with the half-used bottle and frames the next click as a label inspection, the buyer can reconsider L-carnitine without feeling sold the same promise again.",
            14: "Concept inspo: Bathroom cabinet, half-used bottle, label turn, locked vault, amber key, quiet mitochondria gate metaphor, serving-size cards, real JCKED bottle on a counter.",
            15: "Cast type: AI UGC",
            16: "AI avatar look: Male 36 to 42, calm, direct, neutral kitchen or bathroom light, no gym-bro styling, no smiling through the hook.",
            17: "Pacing: medium",
            18: "Main emotion: Recognition first, then a cleaner second chance.",
            19: "Voiceover: Restrained HeyGen narration, recordable as written, no hype cadence.",
            20: "Music: Sparse, quiet, documentary bed at low mix. Drop under the final line.",
            21: "Caption style: Burned-in sans-serif, bottom third, six words or fewer when possible.",
            22: "Overall direction: Editorial science documentary, navy shadows, restrained amber accents, product bottle photographed rather than generated.",
            23: "Editing notes: Open on the cabinet and the old bottle before any biology. Use the vault and key as a metaphor only. Keep production detail out of the Script cells. Product reveal should land on the label and serving size, not a transformation promise.",
            25: "Format: Video",
            27: "Note: Claim guardrails: current PDP scan found JCKED positioning around a 4000MG liquid L-carnitine dose and liquid delivery. Final label art and serving facts still need brand-side verification before send. Safe script use is label inspection and serving comparison. Avoid clinical dose, hard fat-loss claims, age-decline claims, guaranteed results, and unsupported competitor dose contrasts.",
        },
        "rows": [
            {
                "visual": "Open on a quiet bathroom cabinet. A half-used L-carnitine bottle sits behind other daily items.",
                "section": "Hook 1",
                "script": "That half-used L-carnitine bottle in the cabinet may not have been a fair test.",
                "ost": "Check the old bottle",
            },
            {
                "visual": "Same cabinet, then quick cuts of meal prep, gym bag, and the bottle cap untouched.",
                "section": "Hook 2",
                "script": "Same gym. Same meals. But the old bottle still did nothing.",
                "ost": "Same effort, new test",
            },
            {
                "visual": "Close hand turns the old bottle to its serving panel. No science animation yet.",
                "section": "Hook 3",
                "script": "Before judging L-carnitine, turn the bottle around.",
                "ost": "Turn the label",
            },
            {
                "visual": "Hand sets the old bottle down. The frame holds on the unfinished serving panel.",
                "section": "Buyer wound",
                "script": "A lot of men try the ingredient once, feel nothing, and quietly write it off.",
                "ost": "The old test failed",
            },
            {
                "visual": "Cinematic navy mitochondria gate. Amber key approaches. Keep it metaphorical, not textbook.",
                "section": "Mechanism",
                "script": "L-carnitine helps fatty acids reach the mitochondria, where the body turns fuel into work.",
                "ost": "Carnitine as the bridge",
            },
            {
                "visual": "Cut from gate to real JCKED bottle. Slow label turn. Serving card appears beside it.",
                "section": "Proof or label fact",
                "script": "JCKED lists 4,000 milligrams of liquid L-carnitine in a 15 milliliter serving.",
                "ost": "4,000mg per serving",
            },
            {
                "visual": "Old bottle and JCKED bottle on counter. Camera rests on both labels, not a body result.",
                "section": "Product reveal",
                "script": "So the cleaner test is the label, the serving size, and the bottle in front of him.",
                "ost": "Real bottle. Real serving.",
            },
            {
                "visual": "Final product hold in morning light. CTA button appears quietly.",
                "section": "CTA",
                "script": "Compare the serving on the old bottle, then decide if JCKED is worth another look.",
                "ost": "Compare before you decide",
            },
        ],
    },
    "puravita": {
        "out": ROOT / "TrendScale_Puravita_Production_Brief_v10.docx",
        "product": "Puravita",
        "fields": {
            4: "Strategist: Farrice Cain",
            5: "Client: PuraVita Magnesium 12-in-1 Complex Capsules",
            6: "PDP: https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex",
            7: "Media buyer notes: Quiet editorial cold-traffic test. Lead with recognition, not magnesium education. Avoid cure language, certainty about depletion, influencer references, serving-direction claims, and only-formula claims unless cleared.",
            8: "Concept type: Iteration",
            9: "Iterated from: The Battery You Can't See",
            10: "Avatar: Male 38 to 62, sweet spot 45 to 55, recently told his bloodwork looks fine but still feels off, and is looking for an answer before his family starts managing him.",
            11: "Awareness level: Unaware to Problem",
            12: "Hypothesis: If the ad starts with the phone battery metaphor and the normal-bloodwork gap, the buyer will inspect the magnesium form list as a possible missing signal instead of hearing another sleep supplement pitch.",
            14: "Concept inspo: Phone at 5 percent, nightstand at dawn, hand only, normal lab report, warm kitchen, coffee on counter, soft body-map glow, serum window, 12 form cards around real bottle.",
            15: "Cast type: AI UGC",
            16: "AI avatar look: Male 45 to 50, weathered-handsome, gray at temples, soft kitchen light, still hands, no sales smile.",
            17: "Pacing: slow",
            18: "Main emotion: Quiet recognition that fine may not be the whole answer.",
            19: "Voiceover: Restrained HeyGen narration, calm and human, no guru authority drop.",
            20: "Music: Sparse piano, low mix, silence under the final line.",
            21: "Caption style: Burned-in sans-serif, bottom third, six words or fewer when possible.",
            22: "Overall direction: Editorial science-cover feel with deep navy shadows and sage-green highlights. Product bottle photographed, with form-card graphics only as overlays.",
            23: "Editing notes: Start with the phone battery and normal morning scene. Move to the lab-report gap before the body-map metaphor. Keep production direction out of Script cells. Product reveal should be the 12-form label inspection, not a promise of sleep, energy, or recovery.",
            25: "Format: Video",
            27: "Note: Claim guardrails: current PDP scan found PuraVita positioning around a Magnesium 12-in-1 Complex and 12 active magnesium forms. Final label art and Supplement Facts still need brand-side verification before send. Safe script use is form-list inspection and comparison to simpler bottles. Avoid only formula, one-capsule-before-bed, six-to-eight-week result timing, fatigue cure, sleep cure, depletion certainty, and named-influencer claims.",
        },
        "rows": [
            {
                "visual": "Phone at 5 percent on a nightstand at dawn. Hand reaches in. Hand only, never face.",
                "section": "Hook 1",
                "script": "Your phone tells you when it is at five percent. Your body does not.",
                "ost": "No warning light",
            },
            {
                "visual": "Normal lab report on kitchen counter beside untouched coffee.",
                "section": "Hook 2",
                "script": "Bloodwork can look normal while the answer is still hiding.",
                "ost": "Normal is not complete",
            },
            {
                "visual": "Morning kitchen. A man pauses near the coffee. No face. No tired-stock acting.",
                "section": "Hook 3",
                "script": "That fine lab report does not always explain the way the morning feels.",
                "ost": "Fine on paper",
            },
            {
                "visual": "HeyGen avatar in warm kitchen light, direct to camera, still hands.",
                "section": "Buyer wound",
                "script": "He is not shopping for magnesium. He is trying to explain why fine does not feel fine.",
                "ost": "Looking for an answer",
            },
            {
                "visual": "Soft body-map glow. Small bloodstream window, then wider body context. Keep it metaphorical.",
                "section": "Mechanism",
                "script": "Magnesium supports hundreds of enzyme systems, but most of it is stored outside the blood.",
                "ost": "Beyond the blood test",
            },
            {
                "visual": "Real bottle on counter. Twelve simple form cards arrange around it.",
                "section": "Proof or label fact",
                "script": "PuraVita gives the label a wider test: 12 active magnesium forms in one complex.",
                "ost": "12 active forms",
            },
            {
                "visual": "Close on label inspection. Compare to a plain single-form bottle without accusation.",
                "section": "Product reveal",
                "script": "Instead of betting on one form, he can inspect the whole form list on the bottle.",
                "ost": "Inspect the forms",
            },
            {
                "visual": "Final product hold. Low-pressure CTA button. Music drops out.",
                "section": "CTA",
                "script": "Read the label, compare it with the bottle at home, and decide from there.",
                "ost": "Read. Compare. Decide.",
            },
        ],
    },
}


def replace_paragraphs(doc: Document, fields: dict[int, str]) -> None:
    for idx, text in fields.items():
        doc.paragraphs[idx - 1].text = text


def fill_table(doc: Document, rows: list[dict[str, str]]) -> None:
    table = doc.tables[0]
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    while len(table.rows) > len(rows) + 1:
        table._tbl.remove(table.rows[-1]._tr)
    for row_idx, data in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        cells[0].text = data["visual"]
        cells[1].text = data["section"]
        cells[2].text = data["script"]
        cells[3].text = data["ost"]


def make_docx(brief: dict) -> None:
    shutil.copyfile(TEMPLATE, brief["out"])
    doc = Document(str(brief["out"]))
    replace_paragraphs(doc, brief["fields"])
    fill_table(doc, brief["rows"])
    doc.save(str(brief["out"]))


def make_vo_extract() -> None:
    lines = [
        "# TrendScale v10 VO Only Extract",
        "",
        "Source lock: clean rescue packet plus original storyboards, source decks, and template named in the packet.",
        "VO rule: every Script line below is intended as exact spoken paid-ad voiceover.",
        "",
    ]
    for brief in BRIEFS.values():
        lines.append(f"## {brief['product']}")
        lines.append("")
        lines.append(f"Product: {brief['product']}")
        for row in brief["rows"]:
            lines.append(f"Section: {row['section']}")
            lines.append(f"Script: {row['script']}")
            lines.append("")
    (ROOT / "TrendScale_v10_VO_Only_Extract.md").write_text("\n".join(lines), encoding="utf-8")


def make_receipt() -> None:
    receipt = """# TrendScale v10 Quality Receipt

## Source Lock
- Clean rescue packet: `_active/trendscale-brief-revision/TrendScale_Clean_Rescue_Packet.md`
- JCKED storyboard: `/Users/farricecain/Downloads/farrice-trendscale-complete/04_Source_Markdown/JCKED-Storyboard.md`
- Puravita storyboard: `/Users/farricecain/Downloads/farrice-trendscale-complete/04_Source_Markdown/Puravita-Storyboard.md`
- JCKED source deck pages: concept, avatar, script
- Puravita source deck pages: concept, avatar, script
- Template: `/Users/farricecain/Downloads/TrendScale_Master_Brief_Template (1).docx`

## Template Lock
- Preserved 28 document paragraphs.
- Preserved 1 script table.
- Preserved columns: Visual / editing notes, Section, Script, On-screen text.

## Writing Lock
- Script cells are spoken paid-ad VO only.
- Production direction is held in Visual / editing notes and brief fields.
- Claim guardrails are held in Note fields.
- JCKED preserves The Locked Vault.
- Puravita preserves The Battery You Can't See.

## PDP And Label Notes
- JCKED PDP URL used: https://jcked.com/products/liquid-l-carnitine-4000mg-of1
- PuraVita PDP URL used: https://shoppuravita.com/products/puravita%C2%AE-magnesium-complex
- Current PDP facts are treated as product-page facts, with final label art and Supplement Facts verification required before external send.

## Gate Results
- Pending final gate run after file generation.
"""
    (ROOT / "TrendScale_v10_Quality_Receipt.md").write_text(receipt, encoding="utf-8")


def main() -> None:
    for brief in BRIEFS.values():
        make_docx(brief)
    make_vo_extract()
    make_receipt()


if __name__ == "__main__":
    main()
